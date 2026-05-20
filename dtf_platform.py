import os, re, warnings
import numpy as np
import pandas as pd
from classifier import classify as _yaml_classify
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.gridspec import GridSpec
from datetime import datetime
warnings.filterwarnings('ignore')

# ── Config ──
_HERE = os.path.dirname(os.path.abspath(__file__))
DARTMOUTH_API_KEY  = os.getenv('DARTMOUTH_API_KEY', '')
FACE_DIR           = os.getenv('FACE_DIR',           os.path.join(_HERE, 'Case Studies FACE'))
PROCESS_SAFETY_DIR = os.getenv('PROCESS_SAFETY_DIR', os.path.join(_HERE, 'PROCESS SAFETY'))

# ── §1 LLM ──
from langchain_openai import ChatOpenAI
if not DARTMOUTH_API_KEY:
    print('WARNING: DARTMOUTH_API_KEY not set — FPI Analyst chat will be disabled')
    llm = None
else:
    llm = ChatOpenAI(
        model='anthropic.claude-haiku-4-5-20251001',
        api_key=DARTMOUTH_API_KEY,
        base_url='https://chat.dartmouth.edu/api',
        temperature=0.3,
        max_tokens=800,
    )
    print('LLM configured ✓')

# ── §2 Industry Taxonomy ──

# Keywords used to classify FACE case documents AND auto-detect company industry from narratives
INDUSTRY_KEYWORDS = {
    'Construction':          ['construct', 'contractor', 'building', 'scaffold', 'roofer', 'roofing',
                              'concrete', 'excavat', 'trench', 'carpenter', 'ironworker', 'steel erect',
                              'demolit', 'masonry', 'drywall', 'framing', 'foundation', 'crane operator'],
    'Manufacturing':         ['manufactur', 'plant', 'factory', 'production line', 'assembly line',
                              'machine operator', 'press operator', 'foundry', 'stamping', 'injection mold',
                              'conveyor', 'packaging', 'welding shop', 'fabricat'],
    'Oil & Gas':             ['oil well', 'gas well', 'petroleum', 'refinery', 'pipeline', 'drilling',
                              'wellhead', 'oilfield', 'offshore', 'derrick', 'blowout', 'rig ',
                              'natural gas', 'compressor station', 'tank battery'],
    'Mining':                ['mining', 'mine ', 'quarry', 'tunnel', 'ore ', 'coal ', 'surface mine',
                              'underground mine', 'haul road', 'crusher', 'longwall', 'slope mine'],
    'Agriculture':           ['farm', 'agricultur', 'crop', 'livestock', 'tractor', 'grain', 'orchard',
                              'harvest', 'dairy', 'poultry', 'silo', 'auger', 'combine'],
    'Utilities':             ['electric utility', 'power plant', 'lineman', 'line worker',
                              'electric cooperative', 'transmission line', 'substation', 'utility worker',
                              'power line', 'electrical worker', 'energized line'],
    'Transportation':        ['truck driver', 'commercial vehicle', 'semi-truck', 'tractor-trailer',
                              'delivery driver', 'motor carrier', 'long-haul', 'freight'],
    'Warehousing/Logistics': ['warehouse', 'distribution center', 'logistics', 'order picker',
                              'loading dock', 'pallet', 'fulfillment', 'stock room'],
    'Healthcare':            ['hospital', 'healthcare', 'nursing', 'patient', 'medical center',
                              'clinic', 'surgery', 'lab worker'],
    'Forestry':              ['logging', 'logger', 'forestry', 'timber', 'chainsaw', 'felling',
                              'sawmill', 'skidder', 'log truck'],
}

# Which industries are similar enough to use as fallback when exact match has few cases
INDUSTRY_SIMILARITY = {
    'Construction':          ['Mining', 'Oil & Gas', 'Utilities', 'Manufacturing'],
    'Manufacturing':         ['Construction', 'Utilities', 'Warehousing/Logistics'],
    'Oil & Gas':             ['Mining', 'Construction', 'Utilities'],
    'Mining':                ['Oil & Gas', 'Construction', 'Utilities', 'Forestry'],
    'Agriculture':           ['Construction', 'Manufacturing', 'Forestry'],
    'Utilities':             ['Construction', 'Oil & Gas', 'Manufacturing'],
    'Transportation':        ['Warehousing/Logistics', 'Construction'],
    'Warehousing/Logistics': ['Transportation', 'Manufacturing', 'Construction'],
    'Healthcare':            ['Manufacturing'],
    'Forestry':              ['Agriculture', 'Construction', 'Mining'],
}

ALL_INDUSTRIES = sorted(INDUSTRY_KEYWORDS.keys())

def classify_face_industry(text):
    """Assign an industry label to a FACE case document using keyword matching."""
    t = text.lower()
    scores = {ind: sum(1 for kw in kws if kw.lower() in t)
              for ind, kws in INDUSTRY_KEYWORDS.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else 'General'

def detect_company_industry(df):
    """
    Auto-detect the company's industry from incident narrative text.
    Returns (best_industry, score_breakdown_dict).
    """
    narr_col = next((c for c in df.columns if 'narr' in c.lower() or 'desc' in c.lower()), None)
    if not narr_col:
        return None, {}
    text = ' '.join(df[narr_col].fillna('').str.lower().tolist())
    scores = {ind: sum(text.count(kw.lower()) for kw in kws)
              for ind, kws in INDUSTRY_KEYWORDS.items()}
    total = sum(scores.values())
    if total == 0:
        return None, {}
    best = max(scores, key=scores.get)
    breakdown = {k: round(v / total * 100, 1) for k, v in scores.items() if v > 0}
    return best, breakdown

# ── §3 File loaders (only called on first run if index missing) ──
from bs4 import BeautifulSoup
from langchain_core.documents import Document

def load_html(fp):
    with open(fp, 'r', encoding='utf-8', errors='ignore') as f:
        return BeautifulSoup(f.read(), 'html.parser').get_text('\n', strip=True)
def load_txt(fp):
    with open(fp, 'r', encoding='utf-8', errors='ignore') as f: return f.read()
def load_pdf(fp):
    try:
        import pdfplumber
        with pdfplumber.open(fp) as pdf:
            return '\n'.join(p.extract_text() or '' for p in pdf.pages)
    except: return ''
def load_docx(fp):
    try:
        from docx import Document as D
        return '\n'.join(p.text for p in D(fp).paragraphs if p.text.strip())
    except: return ''

# ── §4 Vector Store ──
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

embeddings = HuggingFaceEmbeddings(model_name='all-MiniLM-L6-v2', model_kwargs={'device': 'cpu'})

INDEX_DIR  = os.path.join(_HERE, 'faiss_index')
# Bump version string whenever the index schema changes — triggers a one-time rebuild
FLAG_FILE    = os.path.join(INDEX_DIR, 'build.flag')
FLAG_VERSION = 'v3-process-safety'

needs_build = (not os.path.exists(INDEX_DIR)
               or not os.path.exists(FLAG_FILE)
               or open(FLAG_FILE).read().strip() != FLAG_VERSION)

if not needs_build:
    print('Loading cached FAISS index (v3 — FACE + Process Safety)...')
    try:
        vectorstore = FAISS.load_local(INDEX_DIR, embeddings, allow_dangerous_deserialization=True)
    except Exception as e:
        print(f'WARNING: Could not load FAISS index ({e}) — chat will be disabled')
        vectorstore = None
else:
    print('Building FAISS index — FACE cases + Process Safety documents (one-time, ~2 min)...')
    docs = []

    # ── FACE fatality case studies ──
    print('\n[1/2] Loading FACE case studies...')
    for fn in sorted(os.listdir(FACE_DIR)):
        fp = os.path.join(FACE_DIR, fn)
        if not os.path.isfile(fp): continue
        ext = fn.lower().rsplit('.', 1)[-1] if '.' in fn else ''
        text = ''
        if ext in ('html', 'htm'): text = load_html(fp)
        elif ext == 'txt':         text = load_txt(fp)
        elif ext == 'pdf':         text = load_pdf(fp)
        elif ext == 'docx':        text = load_docx(fp)
        else: continue
        if len(text.strip()) > 50:
            industry = classify_face_industry(text)
            docs.append(Document(page_content=text,
                                 metadata={'source': fn,
                                           'source_type': 'face',
                                           'industry': industry}))
            print(f'  ✓ {fn}  [{industry}]')

    # ── CCPS Beacon process safety bulletins (year subfolders, year ignored) ──
    print('\n[2/2] Loading Process Safety (Beacon) documents...')
    ps_count = 0
    for root, _, files in os.walk(PROCESS_SAFETY_DIR):
        for fn in sorted(files):
            if not fn.lower().endswith('.pdf'): continue
            fp = os.path.join(root, fn)
            text = load_pdf(fp)
            if len(text.strip()) > 50:
                industry = classify_face_industry(text)
                docs.append(Document(page_content=text,
                                     metadata={'source': fn,
                                               'source_type': 'process_safety',
                                               'industry': industry}))
                ps_count += 1
                print(f'  ✓ {fn}  [{industry}]')

    face_count = len(docs) - ps_count
    print(f'\nLoaded {face_count} FACE cases + {ps_count} Process Safety docs = {len(docs)} total ✓')
    splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=100)
    chunks = splitter.split_documents(docs)
    print(f'{len(docs)} docs → {len(chunks)} chunks')
    if not chunks:
        print('WARNING: No documents found — FAISS index will be empty (chat disabled)')
        vectorstore = None
    else:
        vectorstore = FAISS.from_documents(chunks, embeddings)
        vectorstore.save_local(INDEX_DIR)
        with open(FLAG_FILE, 'w') as f:
            f.write(FLAG_VERSION)
        print('FAISS index built and saved ✓')

# Standard retriever kept for the rag_chain (not used in UI, kept for reference)
retriever = vectorstore.as_retriever(search_kwargs={'k': 6}) if vectorstore else None
print('Retriever ready ✓' if retriever else 'Retriever unavailable (no documents)')

# ── §5 Industry-Aware Retrieval ──

def industry_aware_retrieve(query, target_industry, k=6):
    """
    Two-stage retrieval:
      1. Exact industry match — prioritized
      2. Similar industries — used to supplement if exact matches are few
      3. All others — fill remaining slots to always return k results
    Labels are preserved in metadata so the LLM knows which cases are exact vs. similar.
    """
    if vectorstore is None:
        return []
    if not target_industry or target_industry == 'General':
        return [doc for doc, _ in vectorstore.similarity_search_with_score(query, k=k)]

    similar = INDUSTRY_SIMILARITY.get(target_industry, [])
    candidates = vectorstore.similarity_search_with_score(query, k=50)

    exact        = [(doc, s) for doc, s in candidates
                    if doc.metadata.get('industry') == target_industry]
    similar_hits = [(doc, s) for doc, s in candidates
                    if doc.metadata.get('industry') in similar]
    other        = [(doc, s) for doc, s in candidates
                    if doc.metadata.get('industry') not in ([target_industry] + similar)]

    result = exact[:k]
    remaining = k - len(result)
    if remaining > 0:
        result += similar_hits[:remaining]
    remaining = k - len(result)
    if remaining > 0:
        result += other[:remaining]

    return [doc for doc, _ in result[:k]]

def format_docs(docs, target_industry=None):
    """Format retrieved docs, labelling source type, industry match level."""
    parts = []
    for d in docs:
        doc_ind  = d.metadata.get('industry', 'General')
        src_type = d.metadata.get('source_type', 'face')
        prefix   = 'PROCESS SAFETY' if src_type == 'process_safety' else 'FACE'
        if target_industry and doc_ind == target_industry:
            tag = f'{prefix} [{doc_ind} — your industry]'
        elif target_industry and doc_ind in INDUSTRY_SIMILARITY.get(target_industry, []):
            tag = f'{prefix} [{doc_ind} — similar industry]'
        else:
            tag = f'{prefix} [{doc_ind}]'
        parts.append(f'--- {tag}: {d.metadata.get("source","?")} ---\n{d.page_content}')
    return '\n\n'.join(parts)

# ── §6 FPI Framework ──
FPI_SYSTEM_PROMPT = """You are a FPI (Fatality Proximity Index) safety analyst.
You have two evidence sources: FACE fatality case studies and CCPS Process Safety Beacon bulletins.
Use both as evidence. Be concise.
Sources labelled "[your industry]" are from the same industry as the company — weight these highest.
Sources labelled "[similar industry]" are from related industries — use as supporting evidence.
PROCESS SAFETY sources contain detailed lessons learned from process industry incidents — cite them for chemical, pressure, fire, and explosion hazards.

FPI SCALE (1=fatal path active, 5=first aid):
1 Fatal Path: worker in active fatal energy path — little/no margin remaining (SLL, EEE HV, CEE engulf, SME, VPT, EOE direct strike)
2 One Barrier: fatal energy present but one meaningful barrier, distance, or incomplete exposure remains
3 Escalation Needed: injury mechanism exists; fatality requires a major change in conditions or escalation
4 Low Energy: no realistic fatal pathway without multiple additional failures
5 First Aid/Ergonomic: minimal energy; first-aid level or ergonomic/MSD events

FAILURE MODES by energy: Gravitational (falls), Mechanical (caught-in, cave-in),
Kinetic (struck-by), Electrical (shock, arc), Thermal/Chemical (fire, explosion, inhalation),
Atmospheric (confined space, toxic gas, drowning), Pressure, Vehicle/Transport, Ergonomic.

Modifiers toward FPI 1: confined space, no fall protection, trapped, LOTO bypass,
compound hazard (multiple energy types in same incident).

For incidents: state energy type, FPI score, reasoning, FACE parallels, controls.
For org data: identify FPI profile, flag FPI 1-2, flag compound hazards,
recommend by FPI impact not frequency. Flag leading indicators.
Keep answers under 300 words unless asked for detail."""

# ── §7 RAG Chain ──
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import HumanMessage, AIMessage

rag_prompt = ChatPromptTemplate.from_template(
    FPI_SYSTEM_PROMPT + '\n\n## FACE Cases\n{context}\n\n## Query\n{question}')

if llm and retriever:
    rag_chain = (
        {'context': retriever | (lambda docs: format_docs(docs)), 'question': RunnablePassthrough()}
        | rag_prompt | llm | StrOutputParser())
    print('RAG chain ready ✓')
else:
    rag_chain = None
    print('RAG chain skipped (no API key)')

# ── §8 Scoring Engine ──
# Fatal-physics taxonomy — 15 primary categories + HRM as modifier in score_fpi()
# Ordered most-fatal-first so the primary code (first match) reflects highest severity
FML = [
    # ── FET: Fire / Explosion / Thermal Runaway ──────────────────────────────
    (r'explosion|explod(?:ed|ing)?\b|detonat|arc.blast|thermal.runaway|battery.*(?:fire|explod)|'
     r'flammable.*vapor.*ignit|dust.*explosion|flash.*fire|gas.*(?:explosion|ignit)|'
     r'vapor.*cloud.*ignit|bleve|engulfing.*fire|fire.*engulf|caught.*on.*fire|'
     r'clothing.*(?:caught.*fire|ignit)|ignit.*accid|combustion.*accid|structure.*fire.*injur|'
     r'fire.*burn.*(?:worker|employee|person)',
     'FET', 'Fire/Explosion/Thermal Runaway', 'Thermal/Chemical'),

    # ── CEE: Confined / Engulfment / Entrapment ──────────────────────────────
    (r'confined.space|engulf(?:ment|ed)|grain.*(?:bin|engulf)|silo.*(?:entry|engulf)|'
     r'trench.*collaps|excavation.*collaps|cave.?in|buried.*alive|burial|drown(?:ed|ing)|'
     r'submerg|fell.*into.*water|fell.*(?:river|pond|lake|tank|pit)|manhole.*entry|'
     r'tank.*entry|pipe.*entry|tunnel.*collaps|asphyxi|suffoc|oxygen.*defic|'
     r'entered.*confined|worked.*(?:inside|within).*(?:tank|vessel|silo|vault|pit)|'
     r'below.?grade.*entry|valve.*vault.*entry|trench.*entry|excavation.*entry',
     'CEE', 'Confined/Engulfment/Entrapment', 'Atmospheric'),

    # ── SLL: Suspended / Lifted Load Path ────────────────────────────────────
    (r'struck.*(?:crane|rigging|sling|suspended.*load|hoisted|lifted.*(?:material|panel|beam|steel|pipe|load))|'
     r'crane.*(?:load|swing|boom).*(?:struck|hit|fell|drop)|rigging.*fail|sling.*fail|'
     r'load.*drop(?:ped)?|load.*(?:fell|swing|struck)|suspended.*load|'
     r'tag.*line.*(?:struck|snap|whip)|headache.*ball|crane.*block.*fell|'
     r'pick.*and.*carry.*fail|lift.*plan.*fail|'
     r'material.*lifted.*(?:fell|drop|swing)|steel.*erect.*(?:fell|drop)|'
     r'beam.*fall.*while.*lift|below.*crane.*path|under.*suspended|working.*under.*load|'
     r'rigging.*failure|sling.*failure',
     'SLL', 'Suspended/Lifted Load Path', 'Gravitational'),

    # ── SME: Stored Mechanical Energy ────────────────────────────────────────
    (r'cable.*(?:snap|part|broke|fail|whip)|tensioned.*cable|wire.*rope.*(?:snap|fail|part)|'
     r'hydraulic.*(?:burst|hose.*burst|line.*burst|fail|release)|'
     r'pneumatic.*(?:burst|release|hose.*burst)|spring.*(?:release|snap|fail)|'
     r'pressurized.*(?:burst|release|hose|line)|hose.*(?:burst|whip|fail|kicked)|'
     r'stored.*energy.*release|valve.*(?:fail|burst|blow)|fitting.*(?:fail|blow)|'
     r'pressure.*vessel.*fail|accumulator.*fail|torque.*release|winch.*(?:fail|snap)|'
     r'come.*along.*fail|unexpected.*release.*energy|loto.*fail|lockout.*fail|'
     r'pressure.*relief.*fail|unexpected.*energiz|re.?energiz.*accid',
     'SME', 'Stored Mechanical Energy', 'Pressure'),

    # ── LFR: Line-of-Fire Release ─────────────────────────────────────────────
    (r'line.of.fire|recoil|kickback|snap.?back|whip(?:ped|ping)|'
     r'struck.*by.*(?:pipe|hose|cable|wire|chain|rod|bar|beam|board|plank|timber).*(?:flew|swing|whip|release|snap|kick)|'
     r'flying.*(?:debris|object|material|piece|fragment)|ejected.*(?:material|object|part)|'
     r'projectile|spall|ricochet|'
     r'flung.*(?:object|material|piece|debris)|launched.*(?:object|projectile)|'
     r'propelled.*(?:object|material|fragment)|thrown.*(?:object|debris|material)',
     'LFR', 'Line-of-Fire Release', 'Kinetic'),

    # ── EEE: Energized Electrical Exposure ───────────────────────────────────
    (r'electrocution|electrocuted|arc.flash|arc.blast|power.?line.*contact|contact.*power.?line|'
     r'overhead.*line.*contact|energized.*(?:line|conductor|equipment|panel|wire)|high.voltage|'
     r'electr.*(?:shock|contact|burn|fatal)|energized.*contact|induced.*voltage|'
     r'contacted.*(?:electrical|energized)|unintentional.*contact.*electr|'
     r'electrical.*shock|shocked.*by.*(?:electrical|live|energized|wire|current)|'
     r'electrical.*contact|contact.*with.*(?:live.*wire|exposed.*wire|bare.*wire)|'
     r'electric.*current.*(?:pass|flow|through)|ground.*fault.*injur',
     'EEE', 'Energized Electrical Exposure', 'Electrical'),

    # ── EOE: Equipment Operating Envelope ────────────────────────────────────
    (r'struck.*by.*(?:excavator|backhoe|loader|bulldozer|dozer|grader|compactor|roller|paver|'
     r'telehandler|skid.?steer|scraper|crane.*swing|boom.*swing|haul.*truck|dump.*truck|'
     r'water.*truck|concrete.*truck|pump.*truck|man.*lift.*(?:struck|hit))|'
     r'run.*over.*by|run.*down.*by|backed.*over|backing.*(?:struck|hit|over)|'
     r'equipment.*(?:struck|backed|hit).*worker|worker.*struck.*(?:equipment|machine|arm|bucket|boom)|'
     r'equipment.*(?:arm|bucket|blade|drum|attachment).*(?:contact|struck|hit).*worker|'
     r'machine.*(?:arm|bucket|blade).*contact|swing.*(?:arm|bucket|blade).*struck|'
     r'equipment.*backup.*(?:struck|hit)|equipment.*(?:travel|move).*(?:struck|hit)',
     'EOE', 'Equipment Operating Envelope', 'Kinetic'),

    # ── VPT: Vehicle / Public Traffic Interface ───────────────────────────────
    (r'struck.*by.*(?:car|automobile|pickup|semi|tractor.trailer|motorist|passing.*vehicle|third.party.*vehicle)|'
     r'(?:vehicle|car|truck|traffic|motorist).*struck.*worker|roadway.*struck|highway.*struck|'
     r'work.zone.*struck|traffic.*struck|hit.*by.*(?:car|automobile|traffic)|'
     r'flagging.*struck|flagger.*struck|pedestrian.*struck.*(?:vehicle|car|traffic)|'
     r'traffic.*(?:collision|crash).*worker|work.*zone.*(?:intrusion|crash)|'
     r'vehicle.*backup.*struck|struck.*reversing.*vehicle',
     'VPT', 'Vehicle/Traffic Interface', 'Kinetic'),

    # ── MEI: Mobile Equipment Instability ────────────────────────────────────
    (r'overturn|rollover|roll.?over|tipped.*over|tip.?over|'
     r'crane.*(?:tip|overturn|collaps)|telehandler.*(?:tip|overturn|roll)|'
     r'excavator.*(?:tip|overturn|roll)|forklift.*(?:tip|overturn|roll)|'
     r'equipment.*(?:tipped|rolled|overturned|unstable|slid|sank)|'
     r'boom.?truck.*(?:tip|overturn)|aerial.*lift.*(?:tip|overturn|fell)|'
     r'ground.*bearing.*fail|outrigger.*fail|mat.*sink|machine.*overturn|'
     r'crane.*collaps|mast.*collaps|boom.*collaps|jacked.*over|side.?tip|'
     r'lateral.*instability.*equip',
     'MEI', 'Mobile Equipment Instability', 'Kinetic'),

    # ── OGP: Overhead Gravity Path ───────────────────────────────────────────
    (r'struck.*by.*falling|struck.*from.*above|object.*fell.*on|'
     r'struck.*by.*(?:dropped|dislodged|falling)|'
     r'falling.*(?:object|tool|debris|material|bolt|nut|pipe|plank|board|panel|equipment|piece).*(?:struck|hit|fell)|'
     r'(?:tool|bolt|nut|pipe|debris|material|equipment|plank).*fell.*from.*(?:above|overhead|upper|tower|scaffold|roof|structure|floor|crane)|'
     r'dropped.*(?:tool|material|object|equipment).*struck|'
     r'something.*fell.*(?:on|from.*above|overhead)|'
     r'material.*fell.*from.*(?:above|overhead|upper)',
     'OGP', 'Overhead Gravity Path', 'Gravitational'),

    # ── IEA: Improvised Elevated Access ──────────────────────────────────────
    (r'man.*basket.*(?:forklift|telehandler|lift|elevated)|'
     r'(?:forklift|telehandler).*(?:basket|man.*basket|work.*platform)|'
     r'improvised.*(?:platform|scaffold|access)|nonstandard.*(?:platform|access)|'
     r'standing.*on.*(?:forklift|pallet|bucket).*elevated|ladder.*as.*work.*platform|'
     r'working.*from.*(?:pallet|bucket).*elevated|'
     r'non.?standard.*access|unapproved.*(?:platform|access|lift)|'
     r'standing.*on.*(?:vehicle|equipment).*to.*reach|makeshift.*(?:platform|scaffold|access)',
     'IEA', 'Improvised Elevated Access', 'Gravitational'),

    # ── CCZ: Crush / Caught-Between Zone ─────────────────────────────────────
    (r'caught.*between|crush(?:ed|ing)|pinch(?:ed|ing).*between|squeez(?:ed|ing).*between|'
     r'caught.*in.*(?:machine|equipment|gear|conveyor|auger|press|roller|chain|nip.*point)|'
     r'entangled.*in|pulled.*into.*(?:machine|equipment|auger|conveyor|grinder|chipper|stump)|'
     r'drawn.*into|caught.*in.*running|nip.*point|baler.*caught|'
     r'caught.*in.*(?:baler|chipper|grinder|shredder|mixer)|'
     r'arm.*(?:caught|pulled).*into|clothing.*(?:caught|entangled).*in.*(?:machine|equip)|'
     r'pinned.*between|wedged.*between|trapped.*between.*(?:equipment|structure|vehicle)|'
     r'compressed.*between|caught.*in.*rotating',
     'CCZ', 'Crush/Caught-Between Zone', 'Mechanical'),

    # ── HTA: Hazardous Atmosphere / Toxic Exposure ───────────────────────────
    (r'carbon.monoxide\b|co.*poison|hydrogen.sulfide|h2s\b|'
     r'toxic.*(?:gas|fume|atmosphere|vapor)|inhal.*(?:fume|vapor|chemical|solvent|gas|dust)|'
     r'chemical.*expos(?:ure|ed)|vapor.*expos|atmospheric.*hazard|oxygen.*displac|'
     r'chemical.*inhal|methane.*(?:expos|inhal)|ammonia.*(?:expos|inhal)|'
     r'chlorine.*(?:expos|inhal)|chemical.*burn|acid.*(?:expos|spill|contact)|'
     r'caustic.*(?:expos|spill|contact)|solvent.*(?:poison|absorb|inhal)|'
     r'diesel.*fume.*expos|welding.*fume.*expos|toxic.*expos|'
     r'respiratory.*expos.*chem|noxious.*(?:gas|fume|vapor)',
     'HTA', 'Hazardous Atmosphere/Toxic', 'Atmospheric'),

    # ── TFE: Thermal / Environmental Fatal Exposure ───────────────────────────
    (r'heat.?stroke|heat.?death|hypertherm|fatal.*heat|heat.*(?:expos|fatal|death)|'
     r'working.*in.*(?:extreme|excessive).*heat|environmental.*heat|cold.?stress|'
     r'hypotherm|frostbite|cold.*expos.*fatal|wildfire.*smoke|smoke.*inhal|'
     r'fire.*smoke.*inhal|heat.*illness|heat.*exhaustion|heat.*cramp|'
     r'heat.*related.*illness|heat.*(?:fatigue|sick)|'
     r'high.*temperature.*expos|working.*direct.*sun.*ill',
     'TFE', 'Thermal/Environmental Exposure', 'Atmospheric'),

    # ── FG subcodes: Fall Geometry — ordered specific → general ──────────────
    (r'fell.*through.*(?:skylight|hatch|floor.*opening|hole|opening|roof.*opening)|'
     r'fall.*through.*(?:skylight|hatch|floor.*opening|hole|opening)|'
     r'skylight.*(?:fell|broke|gave|collapse)|floor.*opening.*fell|open.*hole.*fell|fell.*in.*hole|'
     r'fall.*through.*(?:floor|ceiling|deck|grating|grate|cover)|'
     r'step.*through.*(?:opening|hole)|broke.*through.*(?:floor|ceiling|roof|skylight|grating)',
     'FG-HOLE', 'Fall Geometry — Open Hole/Skylight', 'Gravitational'),

    (r'aerial.*lift|scissor.*lift|boom.*lift|man.?lift|elevated.*work.*platform|'
     r'jlg\b|genie.*lift|work.*platform.*lift|personnel.*lift|articulating.*lift|'
     r'articulated.*boom|self.?propelled.*lift|self.?propelled.*platform|'
     r'\bawp\b|snorkel.*lift|towable.*lift',
     'FG-LIFT', 'Fall Geometry — Aerial Lift', 'Gravitational'),

    (r'fell.*(?:from|off).*roof|fall.*(?:from|off).*roof|fell.*roof.*edge|'
     r'fell.*from.*(?:peak|ridge|eave|parapet)|roof.*edge.*fall|'
     r'fell.*(?:from|off).*(?:flat.*roof|metal.*roof|pitched.*roof|shingled.*roof)|'
     r'roofing.*fall|working.*on.*roof.*fell|fell.*while.*roofing',
     'FG-ROOF', 'Fall Geometry — Roof Edge', 'Gravitational'),

    (r'fell.*(?:from|off).*scaffold|fall.*(?:from|off).*scaffold|scaffold.*(?:fell|collapse|gave.*way)|'
     r'scaffold.*board.*(?:broke|gave|fell|collapse)|scaffolding.*fell|fell.*scaffolding|'
     r'plank.*(?:gave.*way|broke|slipped).*scaffold|scaffold.*(?:plank|board).*fail|'
     r'working.*scaffold.*fell|fell.*(?:from|off).*scaffolding',
     'FG-SCAF', 'Fall Geometry — Scaffold', 'Gravitational'),

    (r'fell.*(?:from|off).*ladder|fall.*(?:from|off).*ladder|fell.*ladder|'
     r'ladder.*(?:slid|fell|gave|tip)|'
     r'ladder.*(?:slipped|kicked|twisted|rotated|buckled)|'
     r'fell.*(?:while|when).*(?:climbing|descending).*ladder|'
     r'missed.*rung|lost.*footing.*ladder|ladder.*(?:unstable|wobbled|moved).*fell',
     'FG-LAD', 'Fall Geometry — Ladder', 'Gravitational'),

    (r'fell.*from.*(?:tower|structure|building|platform|deck|balcony|mezzanine|catwalk|bridge|walkway|stair|ironwork|steel)|'
     r'fall.*from.*(?:tower|structure|building|platform|deck|mezzanine|catwalk|ironwork|steel)|'
     r'fell.*off.*(?:tower|structure|building|platform|deck|mezzanine|catwalk|ironwork)|'
     r'fell.*from.*(?:work.*platform|elevated.*platform|raised.*platform|work.*deck|loading.*dock)|'
     r'fell.*off.*(?:loading.*dock|dock|elevated.*walkway|overhead.*walkway)|'
     r'fell.*from.*(?:elevated|upper).*level|fell.*while.*(?:ironwork|erection|structural.*work)',
     'FG-STR', 'Fall Geometry — Elevated Structure', 'Gravitational'),

    (r'fell.*from.*(?:truck|vehicle|equipment|machine|trailer|flatbed|step|running.*board|cab)|'
     r'fall.*from.*(?:truck|vehicle|equipment|trailer|flatbed)|'
     r'fell.*off.*(?:truck|trailer|equipment|vehicle|step)|'
     r'fell.*(?:off|from).*(?:back.*of.*truck|bed.*of.*truck|tailgate|truck.*bed)|'
     r'fell.*(?:off|from).*(?:tractor|equipment.*cab|machine.*step|tank.*truck|cement.*truck)|'
     r'dismounting.*(?:truck|vehicle|equipment).*fell|mounting.*(?:truck|vehicle|equipment).*fell',
     'FG-VEH', 'Fall Geometry — Vehicle/Equipment', 'Gravitational'),

    # ── ERG: Ergonomic / MSD ─────────────────────────────────────────────────
    (r'overexertion|repetitive.*motion|ergonomic|musculoskelet|\bmsd\b|'
     r'strain.*(?:back|shoulder|neck|muscle)|sprain.*(?:back|shoulder|wrist|ankle|knee)|'
     r'lift.*injur|awkward.*posture|back.*(?:strain|sprain|injury|pain.*work)|'
     r'shoulder.*(?:strain|sprain|injury)|pulled.*(?:muscle|back)|'
     r'herniated.*disc|rotator.*cuff|carpal.*tunnel|'
     r'cumulative.*trauma|repetitive.*strain|tendinitis|bursitis',
     'ERG', 'Overexertion/MSD', 'Ergonomic'),

    # ── SBO: Struck By — Object (Unspecified) — broad catch-all before FG-GEN ─
    (r'\bstruck\b.*\bby\b|\bhit\b.*\bby\b.*(?:object|bar|pipe|beam|door|gate|post|pole|sign|guard|rail|equipment.*part|machinery.*part)|'
     r'walked.*into.*(?:object|equipment|structure|pipe|bar|beam|pole|column|door|gate|post|guard|rail|wall|vehicle)|'
     r'bumped.*(?:head|into|against)|ran.*into.*(?:object|equipment|structure)|'
     r'came.*into.*contact.*with.*(?:object|structure|equipment|part|surface)|'
     r'worker.*contact.*(?:object|equipment|machinery|part)|'
     r'head.*(?:struck|hit|bumped|knocked).*(?:on|against)|'
     r'body.*part.*(?:struck|hit|contact).*(?:on|against|by)|'
     r'contact.*with.*(?:rotating|moving|sharp|stationary|energized|pressurized).*(?:part|object|surface|equipment|guard|shaft|component|machinery|machine)',
     'SBO', 'Struck By — Object (Unspecified)', 'Kinetic'),

    # ── FG-GEN: Fall Geometry — catch-all (must be last) ────────────────────
    (r'\bfell\b|\bfall(?:en|ing)?\b(?!.*height.*exemp)|\bslipped\b|\btripped\b|'
     r'lost.*(?:footing|balance|grip)|tumbled|pitched.*forward|pitched.*off|'
     r'toppled.*(?:over|off)|knocked.*(?:to.*ground|down.*(?:floor|ground))|'
     r'slid.*(?:off|down|from)|rolled.*(?:off|down)|stumbled.*(?:and|on)|'
     r'lost.*step|missed.*step|went.*down.*(?:floor|ground|stairs)|'
     r'dropped.*to.*(?:ground|floor)|feet.*(?:slipped|gave|went)|'
     r'floor.*(?:gave.*way|collapsed).*fell|ground.*(?:gave.*way|collapsed).*fell',
     'FG-GEN', 'Fall Geometry — General', 'Gravitational'),
]

def classify_fm_multi(evt, narr):
    t = (str(evt) + ' ' + str(narr)).lower()
    matches = []; seen = set()
    for pat, code, label, energy in FML:
        if code not in seen and re.search(pat, t):
            matches.append((code, label, energy)); seen.add(code)
    return matches if matches else [('OTHER', 'Other', 'Other')]

def extract_h(t):
    m = re.search(r'(\d+\.?\d*)\s*(?:foot|feet|ft)\b', str(t).lower())
    return float(m.group(1)) if m else np.nan

def extract_v(t):
    m = re.search(r'(\d+[\d,]*)\s*(?:volt|kv|v\b)', str(t).lower())
    return float(m.group(1).replace(',', '')) if m else np.nan

# Pre-compiled patterns for score_fpi (compiled once at startup, reused per call)
_P = {k: re.compile(v) for k, v in {
    'lfr_partial':   r'deflect|glancing|partial.*path',
    'eee_hv':        r'power.?line|overhead.*line|transmission.*line|high.?voltage|arc.?flash|arc.?blast|electrocution|electrocuted',
    'mei_no_tip':    r'unstable|stabiliz|caught.*self|almost.*tipped|minor.*tip|tilt.*corrected|no.*overturn',
    'ogp_heavy':     r'beam|pipe|panel|steel|lumber|bundle|pallet|heavy|large|concrete|block',
    'ogp_light':     r'small|light|nut|bolt|hand.*tool|pen|marker|chip|splinter',
    'ccz_severe':    r'cave.?in|buried|full.*body|crush.*equip|pinned.*by.*(?:equip|vehicle|structure)|pulled.*into|entangled.*in|drawn.*into',
    'ccz_extremity': r'finger|thumb|toe|foot|hand.*pinch|minor.*pinch',
    'cee_ctrl':      r'entry.*without.*incident|inspection.*entry|routine.*entry|permit.*check.*pass|no.*hazard.*detect|monitoring.*complet|atmospheric.*clear|air.*monitor.*pass',
    'hta_sublethal': r'sub.?lethal|minor.*expos|odor.*detect|monitor.*alarm|low.*level|trace.*level',
    'fet_small':     r'small.*fire|minor.*fire|fire.*extinguish|contained.*fire|fire.*away.*from.*worker',
    'tfe_severe':    r'heat.?stroke|hypertherm|severe.*heat',
    'tfe_mild':      r'mild.*heat|minor.*heat|heat.*rest|heat.*break|slight.*heat',
    'sbo_head':      r'head.*(?:struck|hit|trauma|fracture)|skull.*fracture|unconscious|concussion',
    'sbo_low':       r'minor.*contact|brushed|glancing|low.*speed|slow.*moving',
    'entrapment':    r'trapped|pinned|buried|engulfed|could.*not.*escape|unable.*to.*escape|immobiliz',
    'wpv_severe':    r'shot|gun|firearm|stabbed|stabbing|weapon|strangulat|chok|unconscious|skull|head trauma',
}.items()}


def score_fpi(fm, h, v, narr, all_fms=None):
    """
    Physics-based severity score on a 1–3 scale.

    Scores reflect the energy/force characteristics of the incident mechanism,
    NOT whether the incident resulted in a fatality, injury, or near-miss.
    Two incidents with identical physics must receive the same score regardless
    of their recorded outcome.

    Scale:
      1 = High severity physics — energy/forces with direct fatality potential
          (significant height fall, high-voltage, heavy equipment strike, crushing,
           pressure release, engulfment, acute toxic exposure, explosion)
      2 = Medium severity physics — meaningful injury potential but lower intrinsic
          energy (low-height fall, moderate overhead object, sub-lethal exposure,
           heat illness, minor crush, struck by moderate-mass object)
      3 = Low severity physics — low energy, no credible fatal mechanism
          (ergonomic/MSD, minor contact, mild thermal, low-height slip/trip)
    """
    n = str(narr).lower(); s = 2; why = []   # default: medium physics

    # ── Fall Geometry subcodes ────────────────────────────────────────────────
    if fm == 'FG-HOLE':
        # Any uncontrolled fall through an opening is high-physics (gravity + impact)
        s = 1; why.append('Fall through opening — high gravitational energy')

    elif fm == 'FG-LIFT':
        # Any uncontrolled descent from an aerial lift is fatal-physics regardless of height
        s = 1; why.append('Aerial lift fall — high gravitational energy (always fatal-physics)')

    elif fm == 'FG-ROOF':
        # Roof edge falls: construction roofs are invariably at fatal height; no threshold
        s = 1; why.append('Roof edge fall — high gravitational energy (always fatal-physics)')

    elif fm == 'FG-SCAF':
        # Scaffold falls account for the largest share of construction fall fatalities
        s = 1; why.append('Scaffold fall — high gravitational energy (always fatal-physics)')

    elif fm == 'FG-LAD':
        # Ladder falls: 300+ construction fatalities/year; fatal at any working height
        s = 1; why.append('Ladder fall — high gravitational energy (always fatal-physics)')

    elif fm == 'FG-STR':
        # Falls from elevated structures are inherently fatal-physics — no height floor
        s = 1; why.append('Elevated structure fall — high gravitational energy (always fatal-physics)')

    elif fm == 'FG-VEH':
        # Falls from vehicles/equipment: lower height, medium physics by default
        s = 2; why.append('Fall from vehicle/equipment — medium gravitational energy')
        if pd.notna(h) and h >= 6:
            s = 1; why.append(f'{h:.0f}ft — elevated, high physics')

    elif fm == 'FG-GEN':
        # Generic fall/slip/trip: medium physics; escalate if significant height
        s = 2; why.append('Fall/slip/trip — medium physics')
        if pd.notna(h) and h >= 6:
            s = 1; why.append(f'{h:.0f}ft — significant height, high physics')
        elif not pd.notna(h):
            # No height info — default medium
            pass

    # ── Suspended / Lifted Load Path ─────────────────────────────────────────
    elif fm == 'SLL':
        # Suspended loads under gravitational energy are always high-physics
        s = 1; why.append('Suspended load path — high gravitational + kinetic energy')

    # ── Stored Mechanical Energy ──────────────────────────────────────────────
    elif fm == 'SME':
        # Pressurized/mechanical energy release is high-physics
        s = 1; why.append('Stored energy release — high pressure/kinetic energy')

    # ── Line-of-Fire Release ──────────────────────────────────────────────────
    elif fm == 'LFR':
        # High-velocity objects are high-physics
        s = 1; why.append('Line-of-fire — high kinetic energy projectile/release')
        if _P['lfr_partial'].search(n):
            s = 2; why.append('Partial/deflected path — medium physics')

    # ── Energized Electrical Exposure ─────────────────────────────────────────
    elif fm == 'EEE':
        # High-voltage and direct contact are high-physics
        if _P['eee_hv'].search(n):
            s = 1; why.append('High-voltage/power line — high electrical energy')
        elif pd.notna(v) and v >= 480:
            s = 1; why.append(f'{v:.0f}V — high-voltage, high physics')
        elif pd.notna(v) and v < 50:
            s = 3; why.append(f'{v:.0f}V — extra-low voltage, low physics')
        else:
            s = 2; why.append('Electrical exposure — medium physics')

    # ── Equipment Operating Envelope ──────────────────────────────────────────
    elif fm == 'EOE':
        # Moving heavy equipment carries high mass × velocity — high physics
        s = 1; why.append('Heavy equipment operating envelope — high kinetic energy')

    # ── Vehicle / Traffic Interface ───────────────────────────────────────────
    elif fm == 'VPT':
        # Road vehicles at speed — high kinetic energy
        s = 1; why.append('Vehicle/traffic — high kinetic energy strike')

    # ── Mobile Equipment Instability ──────────────────────────────────────────
    elif fm == 'MEI':
        # Rollover/tip-over of heavy equipment — high physics
        s = 1; why.append('Equipment instability/overturn — high gravitational + crush energy')
        if _P['mei_no_tip'].search(n):
            s = 2; why.append('Instability without overturn — medium physics')

    # ── Overhead Gravity Path ─────────────────────────────────────────────────
    elif fm == 'OGP':
        # In construction, dropped/falling objects are almost always heavy materials;
        # default to high-physics. Only de-escalate when object is explicitly small/light.
        if _P['ogp_light'].search(n):
            s = 3; why.append('Small/light object — low gravitational energy')
        else:
            s = 1; why.append('Overhead gravity — high gravitational energy (construction default)')

    # ── Improvised Elevated Access ────────────────────────────────────────────
    elif fm == 'IEA':
        # Elevated work without proper platform — medium physics by default
        s = 2; why.append('Improvised elevated access — medium physics')
        if pd.notna(h) and h >= 6:
            s = 1; why.append(f'{h:.0f}ft — significant height, high physics')

    # ── Crush / Caught-Between Zone ───────────────────────────────────────────
    elif fm == 'CCZ':
        # Machinery has enormous mechanical advantage — default is high-physics
        # Only de-escalate to medium for isolated extremity pinch with a light component
        if _P['ccz_extremity'].search(n) and not _P['ccz_severe'].search(n):
            s = 2; why.append('Extremity-only pinch with light component — medium mechanical energy')
        else:
            s = 1; why.append('Caught-in/between machinery — high mechanical energy (default fatal-physics)')

    # ── Confined / Engulfment / Entrapment ────────────────────────────────────
    elif fm == 'CEE':
        # Engulfment, drowning, asphyxiation are high-physics atmospheric
        s = 1; why.append('Confined/engulfment — high atmospheric energy hazard')
        if _P['cee_ctrl'].search(n):
            s = 2; why.append('Controlled entry — medium physics (no engulf)')

    # ── Hazardous Atmosphere / Toxic ──────────────────────────────────────────
    elif fm == 'HTA':
        # Acute toxic exposure is high-physics
        s = 1; why.append('Acute toxic/atmospheric exposure — high chemical energy')
        if _P['hta_sublethal'].search(n):
            s = 2; why.append('Sub-lethal concentration — medium physics')

    # ── Fire / Explosion / Thermal Runaway ────────────────────────────────────
    elif fm == 'FET':
        # Explosions and engulfing fire are always high-physics
        s = 1; why.append('Fire/explosion — high thermal/chemical energy')
        if _P['fet_small'].search(n):
            s = 2; why.append('Small/contained fire — medium physics')

    # ── Thermal / Environmental Exposure ──────────────────────────────────────
    elif fm == 'TFE':
        # Severe heat illness: medium physics (gradual thermal loading, not acute)
        s = 2; why.append('Thermal/environmental — medium physics')
        if _P['tfe_severe'].search(n):
            s = 1; why.append('Severe heat exposure — high physiological energy load')
        elif _P['tfe_mild'].search(n):
            s = 3; why.append('Mild thermal exposure — low physics')

    # ── Ergonomic / MSD ───────────────────────────────────────────────────────
    elif fm == 'ERG':
        # Ergonomic events involve low force levels — always low physics
        s = 3; why.append('Ergonomic/MSD — low mechanical energy')

    # ── Struck By — Object (Unspecified) ─────────────────────────────────────
    elif fm == 'SBO':
        # Default medium; escalate for head strikes (high-mass or velocity implied)
        s = 2; why.append('Struck by object — medium kinetic energy')
        if _P['sbo_head'].search(n):
            s = 1; why.append('Head strike — high kinetic energy impact')
        elif _P['sbo_low'].search(n):
            s = 3; why.append('Low-energy contact — low physics')

    # ── Workplace Violence / Physical Assault ─────────────────────────────────
    elif fm == 'WPV':
        # Most physical assaults are medium physics; firearms/weapons escalate to high
        s = 2; why.append('Physical assault — medium kinetic energy')
        if _P['wpv_severe'].search(n):
            s = 1; why.append('Armed/severe assault — high kinetic energy')

    # ── Entrapment modifier (physics signal — not outcome) ────────────────────
    # Being physically trapped indicates ongoing high-energy exposure regardless of outcome
    if _P['entrapment'].search(n):
        if s > 1:
            s = 1; why.append('Physical entrapment — escalated to high physics')

    # ── Compound hazard ────────────────────────────────────────────────────────
    if all_fms and len(all_fms) > 1:
        energies = set(m[2] for m in all_fms if m[0] != 'OTHER')
        if len(energies) > 1:
            s = min(s, 2); why.append('Compound hazard — multiple energy types')

    return s, '; '.join(why)

def process_incident_csv(df, narr_col=None, evt_col=None):
    from concurrent.futures import ThreadPoolExecutor
    import os

    if not narr_col or narr_col not in df.columns:
        narr_col = next((c for c in df.columns if 'narr' in c.lower() or 'desc' in c.lower()), None)
    if not evt_col or evt_col not in df.columns:
        evt_col = next((c for c in df.columns if 'event' in c.lower() and 'title' in c.lower()), None)
    if not narr_col: return None, 'No narrative/description column found. Please select it from the column picker.'
    df = df.copy()
    df['_n'] = df[narr_col].fillna('')
    df['_e'] = df[evt_col].fillna('') if evt_col else ''

    # ── Classification — parallel across rows ──────────────────────────────────
    n_workers = min(8, (os.cpu_count() or 4))
    evts  = df['_e'].tolist()
    narrs = df['_n'].tolist()
    with ThreadPoolExecutor(max_workers=n_workers) as pool:
        results = list(pool.map(lambda t: _yaml_classify(event_title=t[0], narrative=t[1]),
                                zip(evts, narrs)))
    all_matches = pd.Series(results, index=df.index)

    df['FM_Code']       = [r[0][0] for r in results]
    df['FM_Label']      = [r[0][1] for r in results]
    df['FM_Energy']     = [r[0][2] for r in results]
    df['FM_All_Labels'] = [' | '.join(m[1] for m in r) if len(r) > 1 else r[0][1] for r in results]
    df['FM_Compound']   = [len(set(m[2] for m in r if m[0] != 'OTHER')) > 1 for r in results]
    df['FM_Review_Flag'] = df['FM_Code'] == 'OTHER'

    # ── Height / voltage — vectorized pandas extraction ────────────────────────
    df['Height_ft'] = pd.to_numeric(
        df['_n'].str.extract(r'(\d+\.?\d*)\s*(?:foot|feet|ft)\b', flags=re.IGNORECASE, expand=False),
        errors='coerce')
    df['Voltage'] = pd.to_numeric(
        df['_n'].str.extract(r'(\d+[\d,]*)\s*(?:volt|kv)\b', flags=re.IGNORECASE, expand=False
        ).str.replace(',', '', regex=False),
        errors='coerce')

    # ── Scoring — parallel across rows ─────────────────────────────────────────
    codes  = df['FM_Code'].tolist()
    hts    = df['Height_ft'].tolist()
    volts  = df['Voltage'].tolist()
    narr_l = df['_n'].tolist()
    match_l = results

    def _score_row(args):
        fm, h, v, narr, fms = args
        return score_fpi(fm, h, v, narr, all_fms=fms)

    with ThreadPoolExecutor(max_workers=n_workers) as pool:
        scored = list(pool.map(_score_row, zip(codes, hts, volts, narr_l, match_l)))

    df['FPI']     = [s[0] for s in scored]
    df['FPI_Why'] = [s[1] for s in scored]
    df.drop(columns=['_n', '_e'], inplace=True)
    return df, None

# ── §9 Action Plan Generator ──
MITIGATIONS = {
    # New taxonomy labels
    'Fire/Explosion/Thermal Runaway':    ['Hot work permitting + fire watch', 'LEL monitoring during hot work', 'Bonding & grounding', 'Ignition source elimination', 'Emergency egress plan'],
    'Confined/Engulfment/Entrapment':    ['Permit-required confined space program', 'Continuous atmospheric monitoring (O2, LEL, toxics)', 'Attendant + written rescue plan', 'Non-entry rescue equipment'],
    'Suspended/Lifted Load Path':        ['Written lift plan for all critical lifts', 'Radius exclusion zone enforcement', 'Tag lines on all loads', 'Rigger/signalperson certification', 'Load path pre-planning'],
    'Stored Mechanical Energy':          ['Stored energy assessment before work', 'LOTO — de-energize all energy sources', 'Pressure bleed-down before line breaking', 'Hose/cable pre-use inspection'],
    'Line-of-Fire Release':              ['Stand-clear zones during pressurization', 'Line-of-fire hazard pre-task review', 'Inspect hoses/cables for wear before use', 'Pressure relief before disconnecting'],
    'Energized Electrical Exposure':     ['NFPA 70E arc flash risk assessment', 'De-energize + LOTO before work', 'Safe approach distance enforcement', 'Appropriate PPE category (FR/arc-rated)', 'Ground fault protection'],
    'Equipment Operating Envelope':      ['Established exclusion zones around equipment', 'Ground guide / spotter for all backing', 'Proximity alarms / cameras', 'Pre-task communication between operator and workers'],
    'Vehicle/Traffic Interface':         ['Traffic control plan (TCP) for all road work', 'Physical barrier separation from traffic', 'High-visibility PPE for all workers', 'Dedicated flaggers/spotters', 'Speed governors on haul trucks'],
    'Mobile Equipment Instability':      ['Ground bearing capacity assessment before setup', 'Outrigger pad sizing per manufacturer', 'Slope restriction policy', 'ROPS + seatbelt enforcement', 'Anti-tip outrigger interlock'],
    'Overhead Gravity Path':             ['Defined hard hat zones below elevated work', 'Toe boards on all elevated platforms', 'Tool lanyards for all elevated work', 'Debris nets / catch platforms', 'No-drop zone barricading'],
    'Improvised Elevated Access':        ['Prohibit man baskets on forklifts — use aerial lifts', 'Engineered work platforms only', 'Pre-task review of elevation access method', 'Height-specific fall protection'],
    'Crush/Caught-Between Zone':         ['Machine guarding audit + interlocked guards', 'LOTO program for all maintenance', 'Exclusion zones during equipment operation', 'Entanglement hazard JHA before task'],
    'Hazardous Atmosphere/Toxic':        ['Atmospheric testing before and during work', 'Continuous gas detection alarms (CO, H2S, O2)', 'Supplied-air or appropriate APF respirator', 'Ventilation engineering controls'],
    'Thermal/Environmental Exposure':    ['Work/rest cycles + hydration program', 'Heat illness acclimatization schedule', 'Buddy system + supervisor check-ins', 'Cold-stress monitoring in winter operations'],
    'Fall Geometry — Open Hole/Skylight': ['Hole covers — rated, marked, secured', 'Skylight screens / guardrails', 'Permit required for uncovered hole work', 'Fall arrest if cover cannot be maintained'],
    'Fall Geometry — Aerial Lift':       ['Pre-operation inspection checklist', '100% tie-off inside platform at all times', 'Outrigger stability check before elevation', 'No repositioning while elevated', 'Ground spotter'],
    'Fall Geometry — Roof Edge':         ['Guardrails at all unprotected roof edges', 'Personal fall arrest system', 'Warning line system + designated monitor', 'Roof access permit'],
    'Fall Geometry — Scaffold':          ['Competent person inspection before each shift', 'Guardrails on all open sides and ends', 'Proper access ladder / stair tower', 'Scaffold erection by qualified crew'],
    'Fall Geometry — Ladder':            ['3-point contact training and enforcement', 'Ladder inspection program', 'Evaluate alternatives (scaffold, aerial lift)', 'Secure top and base of ladder'],
    'Fall Geometry — Elevated Structure': ['100% tie-off policy on all elevated structures', 'Safety nets below elevated ironwork', 'Controlled access zones at leading edge', 'Fall protection plan before work begins'],
    'Fall Geometry — Vehicle/Equipment': ['3-point contact when mounting/dismounting', 'No jumping from equipment', 'Handholds + anti-slip steps maintained', 'Safe egress procedures in JHA'],
    'Fall Geometry — General':           ['Slip/trip hazard pre-shift inspection', 'Housekeeping program', 'Anti-slip surfaces in wet/oily areas', 'Adequate lighting'],
    'Overexertion/MSD':                  ['Ergonomic assessment for high-risk tasks', 'Mechanical assists / material handling equipment', 'Job rotation program', 'Stretch & flex program'],
    'Struck By — Object (Unspecified)':  ['Pre-task struck-by hazard identification', 'Barricade/exclusion zones around overhead or swing hazards', 'Hard hat + face shield for head/face exposure', 'Keep-clear distances from pinch points and swing paths'],
    'Workplace Violence / Physical Assault': ['Workplace violence prevention program', 'Threat assessment team', 'Security presence / escort protocols', 'De-escalation training for all staff', 'Panic button / emergency alert systems', 'Zero-tolerance violence policy'],
    # Legacy labels (kept for backward compatibility with old scored CSVs)
    'Cave-in / collapse':            ['Trench shoring / shielding', 'Competent person on site', 'Sloping/benching per OSHA 1926 Subpart P'],
    'Confined space':                ['Permit-required confined space program', 'Continuous atmospheric monitoring', 'Attendant + rescue plan'],
    'Toxic atmosphere':              ['Atmospheric testing before entry', 'Supplied-air respirators', 'Ventilation + gas detection alarms'],
    'Explosion':                     ['Hot work permitting', 'LEL monitoring', 'Bonding & grounding', 'Ignition source elimination'],
    'Roadway collision':             ['Dedicated truck routes', 'Spotter + backup cameras', 'High-visibility PPE', 'Speed governors'],
    'Fall >20 ft':                   ['100% tie-off policy', 'Guardrail systems', 'Safety nets', 'Controlled access zones'],
    'Fall from roof':                ['Guardrails at roof edge', 'Personal fall arrest system', 'Warning line + monitor'],
    'Fall from aerial lift':         ['Pre-operation inspection', '100% tie-off inside lift', 'Ground spotter', 'Outrigger stability check'],
    'Caught-in (maintenance)':       ['Lockout/tagout program (LOTO)', 'Machine-specific LOTO procedures', 'Verification before work'],
    'Struck by vehicle':             ['Pedestrian-free zones', 'Spotters for backing', 'Proximity alarms', 'Barrier separation'],
    'Arc flash':                     ['Arc flash risk assessment', 'PPE category labeling', 'De-energize before work', 'Remote racking'],
    'Fire / ignition':               ['Fire watch during hot work', 'Automatic suppression', 'Emergency egress plan'],
    'Vehicle rollover':              ['Seatbelt enforcement', 'Rollover protection (ROPS)', 'Slope assessment'],
    'Caught-in (operating)':         ['Machine guarding audit', 'Interlocked guards', 'E-stop placement review'],
    'Crushing':                      ['Blocking/cribbing procedures', 'Energy isolation', 'Exclusion zones'],
    'Struck by crane/swinging load': ['Lift plan + radius exclusion zone', 'Tag lines on all loads', 'Rigger certification', 'Load path planning'],
    'Struck by falling object':      ['Hard hat zones', 'Toe boards + debris nets', 'Tool lanyards'],
    'Fall 6-20 ft':                  ['Guardrails', 'Personal fall arrest', 'Scaffold inspection program'],
    'Forklift / PIT':                ['Pedestrian separation', 'Operator certification', 'Speed limits + floor markings'],
    'Chemical burn':                 ['SDS review', 'Secondary containment', 'Emergency eyewash/shower'],
    'Chemical inhalation/exposure':  ['Ventilation assessment', 'Air monitoring + exposure limits', 'Supplied-air or APF respirator'],
    'Fall from scaffold':            ['Competent person inspections', 'Guardrails on all open sides', 'Access ladders'],
    'Fall from ladder':              ['3-point contact training', 'Ladder inspection program', 'Consider alternatives (scaffolds, lifts)'],
    'Electrical shock':              ['De-energize + LOTO', 'Insulated tools', 'GFCIs on all circuits'],
    'Heat stress':                   ['Work/rest cycles', 'Hydration program', 'Acclimatization schedule'],
    'Drowning/submersion':           ['Life ring + throw bags at all water edges', 'PFD policy near water', 'Buddy system'],
    'Overexertion/MSD':              ['Ergonomic assessment', 'Mechanical assists / lift equipment', 'Job rotation program'],
}

def get_mitigations(fm_label):
    return MITIGATIONS.get(fm_label, ['Conduct hazard assessment', 'Develop specific SOP', 'Train workers on controls'])

LEADING_INDICATORS = {
    'Fall Geometry — Scaffold':        ('Fall Geometry — Open Hole/Skylight', 0.12, 'Scaffold near-misses precede fatal falls through unprotected openings'),
    'Fall Geometry — Ladder':          ('Fall Geometry — Open Hole/Skylight', 0.09, 'Ladder falls escalate when height or load increases'),
    'Fall Geometry — General':         ('Fall Geometry — Elevated Structure',  0.10, 'Broad-catch falls signal absent fall protection on elevated structures'),
    'Fall Geometry — Aerial Lift':     ('Fall Geometry — Open Hole/Skylight', 0.11, 'Aerial lift near-misses indicate tie-off non-compliance'),
    'Fall Geometry — Vehicle/Equipment': ('Fall Geometry — Elevated Structure', 0.08, 'Truck/equipment falls indicate unsafe egress procedures at height'),
    'Crush/Caught-Between Zone':       ('Equipment Operating Envelope',        0.15, 'Caught-in events precede fatal equipment strikes when exclusion zones absent'),
    'Overhead Gravity Path':           ('Suspended/Lifted Load Path',          0.08, 'Falling objects precede suspended load fatalities when rigging degrades'),
    'Energized Electrical Exposure':   ('Energized Electrical Exposure',       0.11, 'Low-voltage contacts precede high-voltage and arc flash fatalities'),
    'Equipment Operating Envelope':    ('Vehicle/Traffic Interface',           0.13, 'Equipment near-misses consistently precede pedestrian fatalities'),
    'Thermal/Environmental Exposure':  ('Fire/Explosion/Thermal Runaway',      0.07, 'Heat/burn events signal ignition sources near flammable materials'),
    'Hazardous Atmosphere/Toxic':      ('Confined/Engulfment/Entrapment',      0.09, 'Sub-lethal exposures precede fatal confined space events'),
}

def get_leading_indicators(df):
    fpi3_counts = df[df.FPI == 3].FM_Label.value_counts()
    indicators = []
    for fm, cnt in fpi3_counts.items():
        if fm in LEADING_INDICATORS:
            fatal_fm, rate, rationale = LEADING_INDICATORS[fm]
            indicators.append({'fm': fm, 'count': int(cnt), 'leads_to': fatal_fm,
                                'rate': rate, 'projected': max(1, int(cnt * rate)),
                                'rationale': rationale})
    return sorted(indicators, key=lambda x: x['projected'], reverse=True)[:5]

FPI_DESCRIPTIONS = {
    1: 'High-physics fatality mechanism — directly responsible for >30% of occupational fatalities; no meaningful barrier between worker and fatal energy transfer',
    2: 'Credible fatality mechanism — energy present but outcome depends on position, timing, or protective equipment; survivable with variation in conditions',
    3: 'Low-physics mechanism — energy levels cannot directly cause fatality without extraordinary co-factors',
}

def generate_action_plan(df, industry=None, out_path='/tmp/FPI_Action_Plan.png'):
    # ── Palette ──────────────────────────────────────────────────────────────
    WHITE  = '#ffffff'
    GREEN  = '#00693E'   # Dartmouth green
    D_GRAY = '#0f172a'
    M_GRAY = '#64748b'
    L_GRAY = '#f1f5f9'
    BRD    = '#e2e8f0'
    A_YEL  = '#fef3c7'
    A_TXT  = '#92400e'
    FPI_C   = {1:'#dc2626', 2:'#ea580c', 3:'#16a34a'}
    FPI_LBL = {1:'Fatal Path', 2:'One Barrier', 3:'Low Physics'}
    ENERGY_C = {
        'Gravitational':'#3b82f6','Mechanical':'#f59e0b','Kinetic':'#ef4444',
        'Electrical':'#a855f7','Thermal/Chemical':'#f97316','Atmospheric':'#06b6d4',
        'Vehicle/Transport':'#ec4899','Pressure':'#84cc16','Ergonomic':'#22d3ee','Other':'#94a3b8',
    }

    # ── Data ─────────────────────────────────────────────────────────────────
    n       = len(df)
    fpi_d   = {s: int((df.FPI == s).sum()) for s in range(1, 4)}
    n12     = fpi_d[1] + fpi_d[2]
    n_comp  = int(df['FM_Compound'].sum()) if 'FM_Compound' in df.columns else 0
    n_other = int((df.FM_Code == 'OTHER').sum())
    leading = get_leading_indicators(df)

    # Top 15: worst FPI first, then highest count
    top15 = (df.groupby('FM_Label')
               .agg(count  =('FPI', 'count'),
                    min_fpi=('FPI', 'min'),
                    energy =('FM_Energy', lambda x: x.mode().iloc[0]))
               .reset_index()
               .sort_values(['min_fpi', 'count'], ascending=[True, False])
               .head(15)
               .reset_index(drop=True))

    energy_cts = df.FM_Energy.value_counts()
    top_fm     = top15.iloc[0]['FM_Label'] if len(top15) else '–'
    top_energy = energy_cts.index[0] if len(energy_cts) else '–'
    risk_pct   = n12 / n * 100
    risk_lv    = 'HIGH' if risk_pct > 30 else 'MODERATE' if risk_pct > 15 else 'LOW'
    risk_col   = '#dc2626' if risk_lv == 'HIGH' else '#d97706' if risk_lv == 'MODERATE' else '#16a34a'
    ind_str    = industry or 'General'

    # ── Figure & GridSpec ────────────────────────────────────────────────────
    fig = plt.figure(figsize=(16, 36), facecolor=WHITE)
    gs  = GridSpec(5, 2, figure=fig,
                   height_ratios=[1.0, 1.4, 3.2, 14.0, 5.0],
                   hspace=0.03, wspace=0.05,
                   left=0.03, right=0.97, top=0.99, bottom=0.005)

    # ═══════════════════════════════════════════════════════════════════════
    # §A  HEADER
    # ═══════════════════════════════════════════════════════════════════════
    ax_hdr = fig.add_subplot(gs[0, :])
    ax_hdr.set_xlim(0, 1); ax_hdr.set_ylim(0, 1); ax_hdr.axis('off')

    ax_hdr.add_patch(plt.Rectangle((0, 0.72), 1, 0.28,
                                    transform=ax_hdr.transAxes,
                                    facecolor=GREEN, edgecolor='none'))
    ax_hdr.text(0.022, 0.875, 'WORKPLACE SAFETY AWARENESS REPORT',
                color=WHITE, fontsize=16, fontweight='bold',
                transform=ax_hdr.transAxes, va='center', fontfamily='monospace')
    ax_hdr.text(0.980, 0.875, datetime.now().strftime('%B %d, %Y'),
                color='#a7d4b8', fontsize=9.5, transform=ax_hdr.transAxes,
                va='center', ha='right')
    ax_hdr.text(0.022, 0.73,
                'Fatality Proximity Index (FPI) Framework  ·  Applied Energy & Environmental Solutions  ·  ENGM 204',
                color=WHITE, fontsize=8, transform=ax_hdr.transAxes, va='top', alpha=0.80)
    ax_hdr.add_patch(plt.Rectangle((0, 0), 1, 0.72,
                                    transform=ax_hdr.transAxes,
                                    facecolor=L_GRAY, edgecolor=BRD, linewidth=0.8))
    meta = [
        ('INCIDENTS ANALYZED',        f'{n:,}',                       D_GRAY),
        ('INDUSTRY',                   ind_str,                        D_GRAY),
        ('HIGH-SEVERITY ZONE  (FPI 1–2)', f'{n12:,}  ({risk_pct:.1f}%)', '#dc2626'),
        ('RISK LEVEL',                 risk_lv,                        risk_col),
    ]
    for i, (lbl, val, vcol) in enumerate(meta):
        x = 0.022 + i * 0.245
        ax_hdr.text(x, 0.64, lbl, color=M_GRAY, fontsize=7.5, fontweight='bold',
                    transform=ax_hdr.transAxes, va='top', fontfamily='monospace')
        ax_hdr.text(x, 0.26, val, color=vcol, fontsize=14, fontweight='bold',
                    transform=ax_hdr.transAxes, va='top')

    # ═══════════════════════════════════════════════════════════════════════
    # §B  EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════════════════
    ax_sum = fig.add_subplot(gs[1, :])
    ax_sum.set_xlim(0, 1); ax_sum.set_ylim(0, 1); ax_sum.axis('off')
    ax_sum.add_patch(plt.Rectangle((0, 0), 1, 1,
                                    transform=ax_sum.transAxes,
                                    facecolor=WHITE, edgecolor=BRD, linewidth=0.8))
    ax_sum.add_patch(plt.Rectangle((0, 0), 0.005, 1,
                                    transform=ax_sum.transAxes,
                                    facecolor=GREEN, edgecolor='none'))
    ax_sum.text(0.022, 0.90, 'EXECUTIVE SUMMARY',
                color=GREEN, fontsize=9.5, fontweight='bold',
                transform=ax_sum.transAxes, va='top', fontfamily='monospace')
    bullets = [
        f"Risk Level: {risk_lv} — {n12:,} of {n:,} incidents ({risk_pct:.1f}%) fall in the high-severity zone (FPI 1–2).",
        f"Primary Hazard: '{top_fm}' is the most frequent critical failure mode in this dataset.",
        f"Dominant Energy Category: {top_energy} ({int(energy_cts.iloc[0]):,} incidents, {energy_cts.iloc[0]/n*100:.1f}%).",
    ]
    if n_comp > 0:
        bullets.append(f"Compound Hazards: {n_comp} incident(s) involve multiple interacting energy types — highest fatality potential.")
    if n_other > 0:
        bullets.append(f"Unclassified: {n_other} incident(s) require manual review (FM_Review_Flag = True).")
    y0 = 0.74
    for b in bullets:
        ax_sum.text(0.022, y0, f'•   {b}', color=D_GRAY, fontsize=9.0,
                    transform=ax_sum.transAxes, va='top')
        y0 -= 0.185

    # ═══════════════════════════════════════════════════════════════════════
    # §C  FPI SCALE REFERENCE (left)  +  ENERGY BREAKDOWN (right)
    # ═══════════════════════════════════════════════════════════════════════
    ax_fpi = fig.add_subplot(gs[2, 0])
    ax_fpi.set_xlim(0, 1); ax_fpi.set_ylim(0, 1); ax_fpi.axis('off')
    ax_fpi.add_patch(plt.Rectangle((0, 0), 1, 1,
                                    transform=ax_fpi.transAxes,
                                    facecolor=WHITE, edgecolor=BRD, linewidth=0.8))
    ax_fpi.add_patch(plt.Rectangle((0, 0.93), 1, 0.07,
                                    transform=ax_fpi.transAxes,
                                    facecolor=GREEN, alpha=0.10, edgecolor=BRD, linewidth=0.5))
    ax_fpi.text(0.030, 0.965, 'FPI SCALE REFERENCE',
                color=GREEN, fontsize=9, fontweight='bold',
                transform=ax_fpi.transAxes, va='center', fontfamily='monospace')
    row_top  = 0.905
    row_step = 0.270
    for s in range(1, 4):
        y   = row_top - (s - 1) * row_step
        pct = fpi_d[s] / n * 100
        ax_fpi.add_patch(plt.Rectangle((0.022, y - 0.210), 0.018, 0.200,
                                        transform=ax_fpi.transAxes,
                                        facecolor=FPI_C[s], edgecolor='none'))
        ax_fpi.text(0.058, y - 0.018, f'FPI {s}  —  {FPI_LBL[s]}',
                    color=D_GRAY, fontsize=9.5, fontweight='bold',
                    transform=ax_fpi.transAxes, va='top')
        ax_fpi.text(0.058, y - 0.065, FPI_DESCRIPTIONS[s],
                    color=M_GRAY, fontsize=7.8,
                    transform=ax_fpi.transAxes, va='top')
        ax_fpi.text(0.965, y - 0.028, f'{fpi_d[s]:,}\n({pct:.1f}%)',
                    color=FPI_C[s], fontsize=9, fontweight='bold',
                    transform=ax_fpi.transAxes, va='top', ha='right')
        if s < 3:
            ax_fpi.plot([0.022, 0.965], [y - 0.225, y - 0.225],
                        color=BRD, linewidth=0.5, transform=ax_fpi.transAxes)

    ax_nrg = fig.add_subplot(gs[2, 1])
    ax_nrg.set_facecolor(WHITE)
    for sp in ax_nrg.spines.values():
        sp.set_color(BRD); sp.set_linewidth(0.8)
    ax_nrg.spines['top'].set_visible(False)
    ax_nrg.spines['right'].set_visible(False)
    e_colors = [ENERGY_C.get(e, '#94a3b8') for e in energy_cts.index]
    bars_e   = ax_nrg.barh(energy_cts.index, energy_cts.values,
                            color=e_colors, edgecolor=WHITE, linewidth=0.5, height=0.62)
    max_v = int(energy_cts.max()) if len(energy_cts) else 1
    for bar, v in zip(bars_e, energy_cts.values):
        ax_nrg.text(v + max_v * 0.02, bar.get_y() + bar.get_height() / 2,
                    f'{v:,}  ({v / n * 100:.1f}%)',
                    va='center', fontsize=8.5, color=M_GRAY)
    ax_nrg.set_title('ENERGY CATEGORY BREAKDOWN', fontsize=9, color=GREEN,
                      fontweight='bold', loc='left', pad=8, fontfamily='monospace')
    ax_nrg.set_xlabel('Number of incidents', fontsize=8.5, color=M_GRAY, labelpad=4)
    ax_nrg.tick_params(colors=D_GRAY, labelsize=8.5)
    ax_nrg.invert_yaxis()
    ax_nrg.set_xlim(0, max_v * 1.32)

    # ═══════════════════════════════════════════════════════════════════════
    # §D  TOP 15 RISK RANKING TABLE
    # ═══════════════════════════════════════════════════════════════════════
    ax_tbl = fig.add_subplot(gs[3, :])
    ax_tbl.set_xlim(0, 1); ax_tbl.set_ylim(0, 1); ax_tbl.axis('off')
    ax_tbl.add_patch(plt.Rectangle((0, 0), 1, 1,
                                    transform=ax_tbl.transAxes,
                                    facecolor=WHITE, edgecolor=BRD, linewidth=0.8))
    cur_y = 1.0

    # Compound hazard callout banner (conditional)
    if n_comp > 0:
        cb_h = 0.040
        ax_tbl.add_patch(plt.Rectangle((0, cur_y - cb_h), 1, cb_h,
                                        transform=ax_tbl.transAxes,
                                        facecolor=A_YEL, edgecolor='#d97706', linewidth=1.0))
        ax_tbl.text(0.012, cur_y - cb_h / 2,
                    f'  COMPOUND HAZARD ALERT — {n_comp} incident(s) involve multiple interacting energy types. '
                    f'These carry the highest fatality potential and require immediate prioritisation.',
                    color=A_TXT, fontsize=8.5, fontweight='bold',
                    transform=ax_tbl.transAxes, va='center')
        cur_y -= cb_h + 0.004

    # Section title bar
    tb_h = 0.036
    ax_tbl.add_patch(plt.Rectangle((0, cur_y - tb_h), 1, tb_h,
                                    transform=ax_tbl.transAxes,
                                    facecolor=GREEN, edgecolor='none'))
    ax_tbl.text(0.012, cur_y - tb_h / 2,
                'TOP 15 RISK RANKING   ·   Sorted by proximity to fatality (FPI), then by incident frequency',
                color=WHITE, fontsize=9.5, fontweight='bold',
                transform=ax_tbl.transAxes, va='center', fontfamily='monospace')
    cur_y -= tb_h + 0.002

    # Column header bar
    ch_h = 0.026
    ax_tbl.add_patch(plt.Rectangle((0, cur_y - ch_h), 1, ch_h,
                                    transform=ax_tbl.transAxes,
                                    facecolor=L_GRAY, edgecolor=BRD, linewidth=0.4))
    COL_X = {'rank':0.008, 'dtf':0.038, 'fm':0.096, 'count':0.435,
              'energy':0.505, 'controls':0.665}
    for hdr, xc in [('#', COL_X['rank']), ('FPI LEVEL', COL_X['dtf']),
                     ('FAILURE MODE', COL_X['fm']), ('COUNT', COL_X['count']),
                     ('ENERGY TYPE', COL_X['energy']), ('TOP CONTROLS', COL_X['controls'])]:
        ax_tbl.text(xc + 0.005, cur_y - ch_h / 2, hdr,
                    color=M_GRAY, fontsize=7.5, fontweight='bold',
                    transform=ax_tbl.transAxes, va='center', fontfamily='monospace')
    cur_y -= ch_h + 0.001

    # Green divider under column headers
    ax_tbl.plot([0, 1], [cur_y, cur_y], color=GREEN,
                linewidth=1.0, transform=ax_tbl.transAxes, alpha=0.35)

    # Data rows
    n_show  = min(15, len(top15))
    row_h   = cur_y / max(n_show, 1)

    for i, row in top15.iterrows():
        fpi_val = int(row['min_fpi'])
        y_bot   = cur_y - (i + 1) * row_h
        y_mid   = y_bot + row_h / 2

        # Alternating background
        ax_tbl.add_patch(plt.Rectangle((0, y_bot), 1, row_h,
                                        transform=ax_tbl.transAxes,
                                        facecolor=L_GRAY if i % 2 == 0 else WHITE,
                                        edgecolor='none'))

        # Rank
        ax_tbl.text(COL_X['rank'] + 0.010, y_mid, f'{i + 1:02d}',
                    color=M_GRAY, fontsize=9, fontweight='bold',
                    transform=ax_tbl.transAxes, va='center', ha='center')

        # FPI badge (left color strip + text)
        bw = 0.052; bh = row_h * 0.72
        bx = COL_X['dtf'] + 0.002; by = y_mid - bh / 2
        ax_tbl.add_patch(plt.Rectangle((bx, by), bw, bh,
                                        transform=ax_tbl.transAxes,
                                        facecolor=FPI_C[fpi_val], alpha=0.13, edgecolor='none'))
        ax_tbl.add_patch(plt.Rectangle((bx, by), 0.005, bh,
                                        transform=ax_tbl.transAxes,
                                        facecolor=FPI_C[fpi_val], edgecolor='none'))
        ax_tbl.text(bx + 0.010, y_mid + row_h * 0.10, f'FPI {fpi_val}',
                    color=FPI_C[fpi_val], fontsize=8.5, fontweight='bold',
                    transform=ax_tbl.transAxes, va='center')
        ax_tbl.text(bx + 0.010, y_mid - row_h * 0.18, FPI_LBL[fpi_val],
                    color=FPI_C[fpi_val], fontsize=6.8,
                    transform=ax_tbl.transAxes, va='center')

        # Failure mode
        ax_tbl.text(COL_X['fm'] + 0.005, y_mid + row_h * 0.10, row['FM_Label'],
                    color=D_GRAY, fontsize=9, fontweight='bold',
                    transform=ax_tbl.transAxes, va='center')

        # Count
        ax_tbl.text(COL_X['count'] + 0.018, y_mid, f"{int(row['count']):,}",
                    color=D_GRAY, fontsize=9.5, fontweight='bold',
                    transform=ax_tbl.transAxes, va='center', ha='center')

        # Energy
        ax_tbl.text(COL_X['energy'] + 0.005, y_mid, row['energy'],
                    color=M_GRAY, fontsize=8.5,
                    transform=ax_tbl.transAxes, va='center')

        # Top 2 controls
        mits = get_mitigations(row['FM_Label'])
        ax_tbl.text(COL_X['controls'] + 0.005, y_mid + row_h * 0.12,
                    f'→  {mits[0]}',
                    color=GREEN, fontsize=8.0, fontweight='bold',
                    transform=ax_tbl.transAxes, va='center')
        if len(mits) > 1:
            ax_tbl.text(COL_X['controls'] + 0.005, y_mid - row_h * 0.16,
                        f'→  {mits[1]}',
                        color=D_GRAY, fontsize=7.5,
                        transform=ax_tbl.transAxes, va='center')

        # Row separator
        ax_tbl.plot([0, 1], [y_bot, y_bot], color=BRD,
                    linewidth=0.4, transform=ax_tbl.transAxes, alpha=0.9)

    # ═══════════════════════════════════════════════════════════════════════
    # §E  LEADING INDICATORS
    # ═══════════════════════════════════════════════════════════════════════
    ax_lead = fig.add_subplot(gs[4, :])
    ax_lead.set_xlim(0, 1); ax_lead.set_ylim(0, 1); ax_lead.axis('off')
    ax_lead.add_patch(plt.Rectangle((0, 0), 1, 1,
                                     transform=ax_lead.transAxes,
                                     facecolor=WHITE, edgecolor=BRD, linewidth=0.8))
    li_h = 0.14
    ax_lead.add_patch(plt.Rectangle((0, 1 - li_h), 1, li_h,
                                     transform=ax_lead.transAxes,
                                     facecolor='#fffbeb', edgecolor=BRD, linewidth=0.5))
    ax_lead.add_patch(plt.Rectangle((0, 1 - li_h), 0.005, li_h,
                                     transform=ax_lead.transAxes,
                                     facecolor='#d97706', edgecolor='none'))
    ax_lead.text(0.022, 1 - li_h / 2 + 0.022,
                 'LEADING INDICATORS',
                 color='#92400e', fontsize=10, fontweight='bold',
                 transform=ax_lead.transAxes, va='center', fontfamily='monospace')
    ax_lead.text(0.022, 1 - li_h / 2 - 0.022,
                 'FPI-3 patterns present in your data with documented conversion rates to FPI-1 fatalities in FACE case data',
                 color=M_GRAY, fontsize=8.5,
                 transform=ax_lead.transAxes, va='center')

    card_top = 1 - li_h - 0.020
    card_bot = 0.045
    card_h   = card_top - card_bot

    if leading:
        n_cards = len(leading)
        pad     = 0.014
        c_w     = (1.0 - 2 * pad - (n_cards - 1) * pad) / n_cards
        for i, ind in enumerate(leading):
            cx = pad + i * (c_w + pad)
            ax_lead.add_patch(plt.Rectangle((cx, card_bot), c_w, card_h,
                                             transform=ax_lead.transAxes,
                                             facecolor=L_GRAY, edgecolor=BRD, linewidth=0.8))
            ax_lead.add_patch(plt.Rectangle((cx, card_bot + card_h - 0.022), c_w, 0.022,
                                             transform=ax_lead.transAxes,
                                             facecolor='#d97706', edgecolor='none'))
            ix = cx + 0.012
            ty = card_bot + card_h - 0.040
            ax_lead.text(ix, ty,        ind['fm'],
                         color=D_GRAY, fontsize=8, fontweight='bold',
                         transform=ax_lead.transAxes, va='top')
            ax_lead.text(ix, ty - 0.10, f"{ind['count']:,} FPI-3 events",
                         color='#d97706', fontsize=8,
                         transform=ax_lead.transAxes, va='top')
            ax_lead.text(ix, ty - 0.20, f"→  {ind['leads_to']}",
                         color='#dc2626', fontsize=8, fontweight='bold',
                         transform=ax_lead.transAxes, va='top')
            ax_lead.text(ix, ty - 0.30, f"~{ind['rate'] * 100:.0f}% conversion rate",
                         color=M_GRAY, fontsize=7.5,
                         transform=ax_lead.transAxes, va='top')
            ax_lead.text(ix, ty - 0.40, f"Est. {ind['projected']} unrealised FPI-1",
                         color='#dc2626', fontsize=8, fontweight='bold',
                         transform=ax_lead.transAxes, va='top')
            ax_lead.text(ix, ty - 0.52, ind['rationale'],
                         color=M_GRAY, fontsize=7.0,
                         transform=ax_lead.transAxes, va='top')
    else:
        ax_lead.text(0.5, card_bot + card_h / 2,
                     'No FPI-3 leading indicators detected in this dataset.',
                     color=M_GRAY, fontsize=10, transform=ax_lead.transAxes,
                     ha='center', va='center')

    fig.text(0.5, 0.002,
             f'FPI Framework v1.2  ·  Applied Energy & Environmental Solutions  ·  ENGM 204 Group 15  ·  {datetime.now().strftime("%Y-%m-%d")}',
             ha='center', va='bottom', fontsize=7.5, color=M_GRAY)

    fig.savefig(out_path, dpi=150, bbox_inches='tight', facecolor=WHITE, edgecolor='none')
    plt.close(fig)
    return out_path

def generate_html_report(df, industry=None, out_path='/tmp/FPI_Report.html'):
    FPI_C   = {1:'#dc2626', 2:'#ea580c', 3:'#16a34a'}
    FPI_LBL = {1:'Fatal Path', 2:'One Barrier', 3:'Low Physics'}
    FPI_DESC = {
        1: 'Fatal energy active — worker in the fatal path, little or no margin remaining',
        2: 'Fatal energy present — one meaningful barrier or incomplete exposure remains',
        3: 'Low severity physics — no credible fatal mechanism without multiple failures',
    }
    n        = len(df)
    fpi_d    = {s: int((df.FPI == s).sum()) for s in range(1, 4)}
    n12      = fpi_d[1] + fpi_d[2]
    risk_pct = n12 / n * 100
    risk_lv  = 'HIGH' if risk_pct > 30 else 'MODERATE' if risk_pct > 15 else 'LOW'
    risk_col = {'HIGH': '#dc2626', 'MODERATE': '#d97706', 'LOW': '#16a34a'}[risk_lv]
    ind_str  = industry or 'General'

    top5 = (df.groupby('FM_Label')
               .agg(count=('FPI', 'count'), min_fpi=('FPI', 'min'),
                    energy=('FM_Energy', lambda x: x.mode().iloc[0]))
               .reset_index()
               .sort_values(['min_fpi', 'count'], ascending=[True, False])
               .head(5).reset_index(drop=True))
    top_fm  = top5.iloc[0]['FM_Label'] if len(top5) else '—'
    leading = get_leading_indicators(df)

    # ── FPI distribution rows ─────────────────────────────────────────────────
    max_cnt = max(fpi_d.values()) or 1
    fpi_rows = ''
    for s in range(1, 4):
        cnt = fpi_d[s]; pct = cnt / n * 100; bw = int(cnt / max_cnt * 100)
        fpi_rows += (
            f'<tr>'
            f'<td style="width:150px"><span class="badge" style="background:{FPI_C[s]}">'
            f'FPI {s} &mdash; {FPI_LBL[s]}</span></td>'
            f'<td class="desc">{FPI_DESC[s]}</td>'
            f'<td style="width:90px"><div class="bar-wrap"><div class="bar" '
            f'style="width:{bw}%;background:{FPI_C[s]}"></div></div></td>'
            f'<td class="num" style="color:{FPI_C[s]}">{cnt:,}</td>'
            f'<td class="pct">{pct:.1f}%</td>'
            f'</tr>'
        )

    # ── Top 5 risk rows ───────────────────────────────────────────────────────
    risk_rows = ''
    for i, row in top5.iterrows():
        fv   = int(row['min_fpi'])
        mits = get_mitigations(row['FM_Label'])
        ctrls = ''.join(f'<span>{m}</span>' for m in mits[:2])
        risk_rows += (
            f'<tr>'
            f'<td class="rank">{i+1}</td>'
            f'<td><span class="badge" style="background:{FPI_C[fv]}">FPI {fv}</span></td>'
            f'<td><div class="fm">{row["FM_Label"]}</div>'
            f'<div class="fm-sub">{int(row["count"]):,} incidents &middot; {row["energy"]}</div></td>'
            f'<td class="ctrls">{ctrls}</td>'
            f'</tr>'
        )

    # ── Leading indicators ────────────────────────────────────────────────────
    li_html = ''
    if leading:
        cards = ''
        for ind in leading[:3]:
            cards += (
                f'<div class="li-card">'
                f'<div class="li-from">{ind["fm"]}</div>'
                f'<div class="li-arrow">&rarr; escalates to: {ind["leads_to"]}</div>'
                f'<div class="li-cnt">{ind["count"]:,} FPI-3 events in data</div>'
                f'<div class="li-risk">~{ind["projected"]} unrealised FPI-1 events</div>'
                f'</div>'
            )
        li_html = (
            f'<div class="section">'
            f'<h2>Leading Indicators &mdash; Low-Severity Patterns That Escalate to FPI-1</h2>'
            f'<div class="li-grid">{cards}</div>'
            f'</div>'
        )

    css = """
* { box-sizing: border-box; margin: 0; padding: 0; }
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Helvetica, Arial, sans-serif;
       background: #fff; color: #0f172a; font-size: 14px; line-height: 1.5; }
.page { max-width: 960px; margin: 0 auto; }
.header { background: #00693E; color: #fff; padding: 22px 32px; }
.header h1 { font-size: 19px; font-weight: 700; letter-spacing: 0.4px; margin-bottom: 4px; }
.header .sub { font-size: 11px; color: rgba(255,255,255,0.70); }
.kpis { display: grid; grid-template-columns: repeat(4,1fr); border-bottom: 3px solid #00693E; }
.kpi { padding: 16px 22px; border-right: 1px solid #e2e8f0; }
.kpi:last-child { border-right: none; }
.kpi .lbl { font-size: 9.5px; text-transform: uppercase; letter-spacing: 1px;
             color: #64748b; font-weight: 600; margin-bottom: 5px; }
.kpi .val { font-size: 22px; font-weight: 700; line-height: 1.2; }
.section { padding: 22px 32px; border-bottom: 1px solid #e2e8f0; }
.section h2 { font-size: 10px; text-transform: uppercase; letter-spacing: 1.5px;
               color: #00693E; font-weight: 700; margin-bottom: 14px;
               border-left: 3px solid #00693E; padding-left: 8px; }
.fpi-tbl { width: 100%; border-collapse: collapse; }
.fpi-tbl tr { border-bottom: 1px solid #f1f5f9; }
.fpi-tbl tr:last-child { border-bottom: none; }
.fpi-tbl td { padding: 9px 10px; vertical-align: middle; }
.badge { display: inline-block; padding: 3px 9px; border-radius: 4px;
          font-size: 10.5px; font-weight: 700; color: #fff; white-space: nowrap; }
.desc { font-size: 12px; color: #475569; }
.num  { font-size: 14px; font-weight: 700; text-align: right; width: 52px; }
.pct  { font-size: 12px; color: #64748b; text-align: right; width: 48px; }
.bar-wrap { width: 80px; background: #f1f5f9; border-radius: 3px; height: 5px; }
.bar { height: 5px; border-radius: 3px; }
.risk-tbl { width: 100%; border-collapse: collapse; }
.risk-tbl th { font-size: 9.5px; text-transform: uppercase; letter-spacing: 0.8px;
                color: #64748b; font-weight: 600; padding: 7px 10px;
                border-bottom: 2px solid #e2e8f0; text-align: left; background: #f8fafc; }
.risk-tbl td { padding: 11px 10px; border-bottom: 1px solid #f1f5f9; vertical-align: top; }
.risk-tbl tr:last-child td { border-bottom: none; }
.rank { font-size: 13px; font-weight: 700; color: #94a3b8; width: 26px; text-align: center; }
.fm { font-weight: 600; font-size: 13px; margin-bottom: 2px; }
.fm-sub { font-size: 11px; color: #64748b; }
.ctrls { font-size: 11.5px; color: #00693E; }
.ctrls span { display: block; margin-bottom: 3px; }
.ctrls span::before { content: '\\2192  '; font-weight: 700; }
.li-grid { display: grid; grid-template-columns: repeat(3,1fr); gap: 12px; }
.li-card { background: #fffbeb; border: 1px solid #fde68a; border-radius: 6px; padding: 14px; }
.li-from { font-size: 11px; font-weight: 700; color: #92400e; margin-bottom: 5px; }
.li-arrow { font-size: 11px; color: #64748b; margin-bottom: 4px; }
.li-cnt  { font-size: 11px; color: #d97706; margin-bottom: 3px; }
.li-risk { font-size: 12px; color: #dc2626; font-weight: 700; }
.footer { padding: 12px 32px; background: #f8fafc; border-top: 1px solid #e2e8f0;
           font-size: 10px; color: #94a3b8; }
@media print { .page { max-width: 100%; } }
"""

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>FPI Safety Report &mdash; {ind_str}</title>
<style>{css}</style>
</head>
<body>
<div class="page">
  <div class="header">
    <h1>FATALITY PROXIMITY INDEX &mdash; SAFETY REPORT</h1>
    <div class="sub">{ind_str} &middot; {datetime.now().strftime('%B %d, %Y')} &middot; Applied Energy &amp; Environmental Solutions &middot; ENGM 204</div>
  </div>
  <div class="kpis">
    <div class="kpi"><div class="lbl">Incidents Analyzed</div><div class="val">{n:,}</div></div>
    <div class="kpi"><div class="lbl">High-Severity Zone (FPI 1&ndash;2)</div>
      <div class="val" style="color:{risk_col}">{n12:,} <span style="font-size:14px">({risk_pct:.1f}%)</span></div></div>
    <div class="kpi"><div class="lbl">Risk Level</div>
      <div class="val" style="color:{risk_col}">{risk_lv}</div></div>
    <div class="kpi"><div class="lbl">Top Failure Mode</div>
      <div class="val" style="font-size:14px;font-weight:600;line-height:1.3">{top_fm}</div></div>
  </div>
  <div class="section">
    <h2>FPI Distribution</h2>
    <table class="fpi-tbl"><tbody>{fpi_rows}</tbody></table>
  </div>
  <div class="section">
    <h2>Top 5 Critical Risks</h2>
    <table class="risk-tbl">
      <thead><tr><th>#</th><th>FPI</th><th>Failure Mode</th><th>Top Controls</th></tr></thead>
      <tbody>{risk_rows}</tbody>
    </table>
  </div>
  {li_html}
  <div class="footer">FPI Framework v1.2 &middot; Applied Energy &amp; Environmental Solutions &middot; ENGM 204 Group 15 &middot; {datetime.now().strftime('%Y-%m-%d')}</div>
</div>
</body>
</html>"""

    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(html)
    return html, out_path

def build_org_summary(df, industry=None):
    n = len(df)
    fpi_d = {s: int((df.FPI == s).sum()) for s in range(1, 4)}
    top_fm = df.FM_Label.value_counts().head(5)
    d1_fm  = df[df.FPI == 1].FM_Label.value_counts().head(3) if (df.FPI == 1).any() else pd.Series(dtype=int)
    n_other    = int((df.FM_Code == 'OTHER').sum())
    n_compound = int(df['FM_Compound'].sum()) if 'FM_Compound' in df.columns else 0

    industry_line = f'Industry: {industry}\n' if industry else ''

    dept_col = next((c for c in df.columns
                     if any(k in c.lower() for k in ['dept', 'department', 'division', 'group', 'site', 'location'])), None)
    dept_info = ''
    if dept_col:
        top_dept = df.groupby(dept_col)['FPI'].agg(['mean', 'count']).sort_values('mean').head(5)
        dept_info = f'\nDepartments by avg FPI (lower = more severe):\n{top_dept.to_string()}'

    sev_col = next((c for c in df.columns
                    if any(k in c.lower() for k in ['severity', 'recordable', 'classification', 'type', 'category'])), None)
    sev_info = ''
    if sev_col:
        sev_cts = df[sev_col].value_counts().head(5)
        sev_info = f'\nIncident classifications:\n{sev_cts.to_string()}'

    similar = INDUSTRY_SIMILARITY.get(industry, [])
    face_context = ''
    if industry:
        face_context = (f'\nFACE retrieval: prioritizing [{industry}] cases'
                        + (f', supplemented by [{", ".join(similar[:2])}] cases' if similar else ''))

    return f"""{industry_line}ORG DATA ({n:,} records)
FPI 1 (Fatal path active):     {fpi_d[1]:,} ({fpi_d[1]/n*100:.1f}%)
FPI 2 (One barrier remaining): {fpi_d[2]:,} ({fpi_d[2]/n*100:.1f}%)
FPI 3 (Low physics severity):  {fpi_d[3]:,} ({fpi_d[3]/n*100:.1f}%)
High-severity zone (FPI 1-2): {fpi_d[1]+fpi_d[2]:,} ({(fpi_d[1]+fpi_d[2])/n*100:.1f}%)
Compound hazard incidents: {n_compound:,}
Unclassified (needs review): {n_other:,}
Top failure modes:
{top_fm.to_string()}
FPI-1 failure modes: {d1_fm.to_string() if len(d1_fm) > 0 else 'None'}{dept_info}{sev_info}{face_context}"""

def filter_by_industry(df, industry):
    """Return rows whose narrative text contains at least one keyword for the given industry."""
    keywords = INDUSTRY_KEYWORDS.get(industry, [])
    if not keywords:
        return df.iloc[0:0]  # empty frame with same columns
    narr_col = next((c for c in df.columns if 'narr' in c.lower() or 'desc' in c.lower()), None)
    if not narr_col:
        return df.iloc[0:0]
    pattern = '|'.join(re.escape(kw) for kw in keywords)
    mask = df[narr_col].fillna('').str.lower().str.contains(pattern, regex=True)
    return df[mask].copy()

print('Scoring engine ready ✓')

# ── §10 Gradio Interface ──
import gradio as gr

current_org = {'summary': '', 'df': None, 'df_industry': None, 'industry': None}

def _fpi_table_lines(df, LBL):
    """Return markdown lines for a FPI breakdown table + near-fatal callout."""
    lines = ['| FPI | Level | Count | % |', '|:---:|:---|---:|---:|']
    for s in range(1, 4):
        cnt = int((df.FPI == s).sum())
        lines.append(f'| {s} | {LBL[s]} | {cnt:,} | {cnt/len(df)*100:.1f}% |')
    n12 = int((df.FPI <= 2).sum())
    lines.append(f'\n**{n12:,} incidents ({n12/len(df)*100:.1f}%) in HIGH-SEVERITY ZONE (FPI 1–2)**')
    top3 = df.FM_Label.value_counts().head(3)
    lines.append('Top failure modes: ' + ' · '.join(f'{lbl} ({cnt})' for lbl, cnt in top3.items()))
    return lines

def upload_and_score(file, industry_selection, narr_col_sel, evt_col_sel):
    if file is None: return 'Upload a file.', None, None, None, None
    fn = file.name.lower()
    try:
        if fn.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file.name)
        else:
            df = pd.read_csv(file.name, low_memory=False)
    except Exception as e:
        return f'Error reading file: {e}', None, None, None, None
    scored, err = process_incident_csv(df,
                                       narr_col=narr_col_sel or None,
                                       evt_col=evt_col_sel or None)
    if err: return f'Error: {err}', None, None, None, None

    # Resolve industry: auto-detect or use selection
    user_picked = industry_selection != 'Auto-detect from my data'
    if user_picked:
        industry = industry_selection
        breakdown = {}
    else:
        detected, breakdown = detect_company_industry(scored)
        industry = detected

    current_org['df']          = scored
    current_org['df_industry'] = None
    current_org['industry']    = industry
    current_org['summary']     = build_org_summary(scored, industry)

    LBL = {1: 'Fatal Path', 2: 'One Barrier', 3: 'Low Physics'}

    # ── §A  OVERALL overview ─────────────────────────────────────────────────
    lines = [f'## Overall — {len(scored):,} incidents scored\n']
    if industry:
        similar = INDUSTRY_SIMILARITY.get(industry, [])
        src = 'Selected' if user_picked else 'Auto-detected'
        lines.append(f'**Industry ({src}): {industry}**')
        if breakdown:
            top3b = sorted(breakdown.items(), key=lambda x: x[1], reverse=True)[:3]
            lines.append('Signal breakdown: ' + ' · '.join(f'{k} {v}%' for k, v in top3b))
        if similar:
            lines.append(f'FACE: prioritizing **{industry}** · fallback **{", ".join(similar[:2])}**')
    else:
        lines.append('*Industry not detected — using full FACE library.*')
    lines.append('')
    lines.extend(_fpi_table_lines(scored, LBL))

    n_other = int((scored.FM_Code == 'OTHER').sum())
    if n_other > 0:
        lines.append(f'\n> ⚠️ **{n_other:,} unclassified** — check `FM_Review_Flag = True` rows.')
    n_compound = int(scored['FM_Compound'].sum()) if 'FM_Compound' in scored.columns else 0
    if n_compound > 0:
        lines.append(f'\n> 🔴 **{n_compound:,} compound hazard incidents** (FPI auto-elevated).')

    # ── §B  INDUSTRY-SPECIFIC overview ───────────────────────────────────────
    ind_df = None
    industry_csv_path  = None
    chart_industry_path = None

    if user_picked and industry:
        ind_df = filter_by_industry(scored, industry)
        if len(ind_df) > 0:
            current_org['df_industry'] = ind_df
            industry_csv_path = '/tmp/fpi_scored_industry.csv'
            ind_df.to_csv(industry_csv_path, index=False)

            lines.append(f'\n---\n## {industry} Subset — {len(ind_df):,} of {len(scored):,} incidents\n')
            lines.extend(_fpi_table_lines(ind_df, LBL))

            n_oi = int((ind_df.FM_Code == 'OTHER').sum())
            if n_oi > 0:
                lines.append(f'\n> ⚠️ **{n_oi:,} unclassified** in this subset.')
            n_ci = int(ind_df['FM_Compound'].sum()) if 'FM_Compound' in ind_df.columns else 0
            if n_ci > 0:
                lines.append(f'\n> 🔴 **{n_ci:,} compound hazard incidents** in this subset.')
        else:
            lines.append(
                f'\n> ℹ️ No incidents matched **{industry}** keywords — industry-specific outputs not generated.'
            )

    # ── CSVs ────────────────────────────────────────────────────────────────
    csv_path = '/tmp/fpi_scored_all.csv'
    scored.to_csv(csv_path, index=False)

    # ── HTML reports ──────────────────────────────────────────────────────────
    _, report_all_path = generate_html_report(scored, industry,
                                              out_path='/tmp/FPI_Report_All.html')
    report_ind_path = None
    if ind_df is not None and len(ind_df) > 0:
        _, report_ind_path = generate_html_report(ind_df, industry,
                                                  out_path='/tmp/FPI_Report_Industry.html')

    return '\n'.join(lines), csv_path, industry_csv_path, report_all_path, report_ind_path

qa_prompt = ChatPromptTemplate.from_messages([
    ('system', FPI_SYSTEM_PROMPT + '\n\n## FACE Cases:\n{context}\n\n## Org Data:\n{org_data}'),
    MessagesPlaceholder(variable_name='chat_history'),
    ('human', '{input}'),
])

def retrieve_and_build(inp):
    target_industry = current_org.get('industry')
    docs = industry_aware_retrieve(inp['input'], target_industry, k=6)
    return {
        'context':      format_docs(docs, target_industry),
        'input':        inp['input'],
        'chat_history': inp['chat_history'],
        'org_data':     current_org.get('summary', 'No org data uploaded yet.'),
    }

if llm:
    convo_chain = RunnableLambda(retrieve_and_build) | qa_prompt | llm | StrOutputParser()
else:
    convo_chain = None

def chat_fn(msg, hist):
    if not convo_chain:
        return 'FPI Analyst chat is not available — DARTMOUTH_API_KEY environment variable is not set.'
    lc = []
    for turn in hist:
        if isinstance(turn, (list, tuple)) and len(turn) == 2:
            u, b = turn
            lc.append(HumanMessage(content=u))
            if b: lc.append(AIMessage(content=b))
    return convo_chain.invoke({'input': msg, 'chat_history': lc})

def on_file_change(file):
    """Auto-detect columns when a file is uploaded; populate column picker dropdowns."""
    empty = gr.Dropdown(choices=[], value=None)
    if file is None:
        return empty, empty
    try:
        fn = file.name.lower()
        if fn.endswith(('.xlsx', '.xls')):
            tmp = pd.read_excel(file.name, nrows=0)
        else:
            tmp = pd.read_csv(file.name, nrows=0)
        cols      = list(tmp.columns)
        narr_auto = next((c for c in cols if 'narr' in c.lower() or 'desc' in c.lower()), None)
        evt_auto  = next((c for c in cols if 'event' in c.lower() and 'title' in c.lower()), None)
        return gr.Dropdown(choices=cols, value=narr_auto), gr.Dropdown(choices=cols, value=evt_auto)
    except Exception:
        return empty, empty

with gr.Blocks(title='FPI Safety Intelligence', theme=gr.themes.Soft()) as demo:
    gr.Markdown(
        '# FPI Safety Intelligence Platform\n'
        '### Fatality Proximity Index — Beyond Safe vs Unsafe\n\n'
        'Upload incident data → Get FPI scores → Download report → Chat with analyst\n\n'
        '*FPI 1 = fatal energy active · FPI 2 = one barrier remaining · FPI 3 = low physics severity*'
    )
    with gr.Tabs():
        # ── Tab 1: Upload & Score ─────────────────────────────────────────────
        with gr.TabItem('Upload & Score'):
            with gr.Row():
                industry_dd = gr.Dropdown(
                    choices=['Auto-detect from my data'] + ALL_INDUSTRIES,
                    value='Auto-detect from my data',
                    label='Industry',
                    info='Auto-detect reads your narratives and picks the closest match.'
                )
            file_in = gr.File(
                label='Upload Incident File (CSV or Excel)',
                file_types=['.csv', '.xlsx', '.xls']
            )
            with gr.Row():
                narr_col_dd = gr.Dropdown(
                    label='Narrative / Description column  (required)',
                    info='Auto-filled on upload — change if wrong.',
                    interactive=True
                )
                evt_col_dd = gr.Dropdown(
                    label='Event title column  (optional)',
                    info='Auto-filled on upload — leave blank if not present.',
                    interactive=True
                )
            file_in.change(on_file_change,
                           inputs=[file_in],
                           outputs=[narr_col_dd, evt_col_dd])
            score_btn = gr.Button('Score All Incidents', variant='primary')
            results   = gr.Markdown()
            with gr.Row():
                csv_dl          = gr.File(label='Download All Incidents (Scored CSV)')
                csv_industry_dl = gr.File(label='Download Industry Subset (Scored CSV)')
            with gr.Row():
                report_dl     = gr.File(label='Download Report — All Incidents (HTML)')
                report_ind_dl = gr.File(label='Download Report — Industry Subset (HTML)')
            score_btn.click(
                upload_and_score,
                inputs=[file_in, industry_dd, narr_col_dd, evt_col_dd],
                outputs=[results, csv_dl, csv_industry_dl, report_dl, report_ind_dl]
            )

        # ── Tab 2: FPI Report ─────────────────────────────────────────────────
        with gr.TabItem('FPI Report'):
            gr.Markdown('### FPI Safety Report\nAfter uploading data in Tab 1, click to regenerate.')
            regen_btn = gr.Button('Regenerate Reports', variant='primary')
            with gr.Row():
                plan_html     = gr.HTML(label='All Incidents — Report Preview')
                plan_html_ind = gr.HTML(label='Industry Subset — Report Preview')
            with gr.Row():
                plan_dl     = gr.File(label='Download — All Incidents (HTML)')
                plan_dl_ind = gr.File(label='Download — Industry Subset (HTML)')
            def regen():
                if current_org['df'] is None: return '<p>No data loaded yet.</p>', '', None, None
                html_all, path_all = generate_html_report(
                    current_org['df'], current_org.get('industry'),
                    out_path='/tmp/FPI_Report_All.html')
                ind_df = current_org.get('df_industry')
                if ind_df is not None and len(ind_df) > 0:
                    html_ind, path_ind = generate_html_report(
                        ind_df, current_org.get('industry'),
                        out_path='/tmp/FPI_Report_Industry.html')
                else:
                    html_ind, path_ind = '<p>No industry subset loaded.</p>', None
                return html_all, html_ind, path_all, path_ind
            regen_btn.click(regen, outputs=[plan_html, plan_html_ind, plan_dl, plan_dl_ind])

        # ── Tab 3: Analyst Chat ───────────────────────────────────────────────
        with gr.TabItem('FPI Analyst Chat'):
            gr.Markdown('Ask about your data, specific incidents, or safety strategy.')
            chatbot    = gr.Chatbot(height=400)
            chat_input = gr.Textbox(placeholder='Ask a safety question...', label='Your question')
            chat_btn   = gr.Button('Send', variant='primary')
            def respond(msg, history):
                if not msg.strip(): return '', history
                history = history + [[msg, chat_fn(msg, history)]]
                return '', history
            chat_btn.click(respond, inputs=[chat_input, chatbot], outputs=[chat_input, chatbot])
            chat_input.submit(respond, inputs=[chat_input, chatbot], outputs=[chat_input, chatbot])

    gr.Markdown('---\n*FPI v1.2 · Applied Energy & Environmental Solutions · ENGM 204 Group 15*')

if __name__ == '__main__':
    demo.launch(debug=True, share=True)
